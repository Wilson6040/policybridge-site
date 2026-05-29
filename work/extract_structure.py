import sys
from docx import Document
from docx.oxml.ns import qn

def describe(path, label, max_paras=99999):
    doc = Document(path)
    print(f"\n{'='*80}\n{label}: {path}\n{'='*80}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    print(f"Total sections: {len(doc.sections)}")
    # Section page setup
    for i, s in enumerate(doc.sections):
        try:
            cols = s._sectPr.find(qn('w:cols'))
            numcols = cols.get(qn('w:num')) if cols is not None else None
        except Exception:
            numcols = None
        print(f"  Section {i}: type={s.start_type}, cols={numcols}, "
              f"page_w={s.page_width}, page_h={s.page_height}")
    # Style usage count
    from collections import Counter
    styles = Counter()
    for p in doc.paragraphs:
        styles[p.style.name] += 1
    print("Paragraph style usage:")
    for st, c in styles.most_common(40):
        print(f"   {st}: {c}")

if __name__ == '__main__':
    describe('/app/work/source/OLD_0223C.docx', 'OLD 0223C')
    describe('/app/work/source/NEW_0526.docx', 'NEW 0526')
    describe('/app/work/source/TEMPLATE_SummaryOfCover.docx', 'TEMPLATE Summary of Cover')
