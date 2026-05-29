from docx import Document
from docx.oxml.ns import qn

def para_text(p):
    return p.text

def dump_headings(path, outpath):
    doc = Document(path)
    lines = []
    for i, p in enumerate(doc.paragraphs):
        st = p.style.name
        txt = p.text.strip()
        if st.startswith('Heading') or st in ('Title','Subtitle'):
            lines.append(f"[{i}] <{st}> {txt}")
    with open(outpath,'w') as f:
        f.write('\n'.join(lines))
    print(f"{path}: {len(lines)} headings -> {outpath}")

def dump_all(path, outpath, limit=None):
    doc = Document(path)
    lines = []
    for i, p in enumerate(doc.paragraphs):
        st = p.style.name
        txt = p.text
        lines.append(f"[{i}|{st}] {txt}")
        if limit and i>=limit:
            break
    with open(outpath,'w') as f:
        f.write('\n'.join(lines))
    print(f"{path}: dumped {len(lines)} paras -> {outpath}")

dump_headings('/app/work/source/OLD_0223C.docx', '/app/work/old_headings.txt')
dump_headings('/app/work/source/NEW_0526.docx', '/app/work/new_headings.txt')
dump_all('/app/work/source/NEW_0526.docx', '/app/work/new_toc_region.txt', limit=400)
dump_all('/app/work/source/TEMPLATE_SummaryOfCover.docx', '/app/work/template_full.txt')
