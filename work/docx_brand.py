"""Shared branding helpers for building TMHCC branded Word documents
from the supplied Summary-of-Cover template (design source)."""
from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---- Brand palette (extracted from template theme) ----
BLUE   = "009CE5"   # accent1 - primary brand blue
GOLD   = "C79000"   # accent2 - subtitle gold / amber
DBLUE  = "0066CC"   # accent3
RED    = "E20033"   # accent4
ORANGE = "FF9900"   # accent5
GREY   = "595959"
DARK   = "212121"

# Traffic-light tints (light background + dark text -> always readable)
TL = {
    "favourable": ("E2F3E8", "1B5E20", "Favourable"),
    "broader":    ("E2F3E8", "1B5E20", "Broader"),
    "neutral":    ("FBF1D6", "7A5800", "Neutral"),
    "clarify":    ("FBF1D6", "7A5800", "Clarification"),
    "restrictive":("FBE3E8", "A4001F", "Restrictive"),
    "new":        ("E2F3E8", "1B5E20", "New cover"),
    "removed":    ("FBE3E8", "A4001F", "Removed"),
}

TEMPLATE = '/app/work/source/TEMPLATE_SummaryOfCover.docx'


def load_template():
    return Document(TEMPLATE)


def clear_body_after_cover(doc, keep_para_index=5):
    """Keep the cover (date table + title + subtitles + wave graphic) and the
    trailing body sectPr; remove everything else so we can rebuild the body."""
    body = doc.element.body
    keep_until = doc.paragraphs[keep_para_index]._p
    removing = False
    for child in list(body):
        if removing:
            if child.tag == qn('w:sectPr'):
                continue
            body.remove(child)
        if child is keep_until:
            removing = True
    return doc


def _style_run(run, size=10, bold=False, color=None, italic=False, font="Arial"):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), font)


def add_para(doc, text="", size=10, bold=False, color=None, italic=False,
             space_after=6, space_before=0, align=None, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        r = p.add_run(text)
        _style_run(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_h1(doc, text):
    """Section-style heading (blue, like template SECTION headings)."""
    p = doc.add_paragraph(style='Heading 1')
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    _style_run(r, size=15, bold=True, color=BLUE)
    return p


def add_h2(doc, text, color=None):
    p = doc.add_paragraph(style='Heading 2')
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    _style_run(r, size=11, bold=True, color=color or BLUE)
    return p


def add_bullet(doc, text, level=0, color=None, bold=False):
    p = doc.add_paragraph(style='List Bullet')
    if level:
        p.paragraph_format.left_indent = Pt(18 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    _style_run(r, size=10, bold=bold, color=color)
    return p


def _set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def _set_cell_text(cell, text, size=9, bold=False, color=None, align=None, bg=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parts = text.split("\n")
    for i, part in enumerate(parts):
        if i:
            p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(1)
        r = p.add_run(part)
        _style_run(r, size=size, bold=bold, color=color)
    if bg:
        _set_cell_bg(cell, bg)


def _set_table_borders(table, color="BFBFBF", sz=4):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        borders.append(e)
    tblPr.append(borders)


def _set_table_full_width(table):
    tblPr = table._tbl.tblPr
    w = OxmlElement('w:tblW')
    w.set(qn('w:type'), 'pct'); w.set(qn('w:w'), '5000')
    tblPr.append(w)
    table.autofit = False


def add_table(doc, headers, rows, widths=None, status_col=None, header_bg=BLUE,
              body_size=9, header_size=9):
    """headers: list[str]; rows: list[list[str]].
    status_col: index of column whose value is a TL key -> coloured chip."""
    ncols = len(headers)
    table = doc.add_table(rows=1, cols=ncols)
    _set_table_borders(table)
    _set_table_full_width(table)
    # header
    for j, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[j], h, size=header_size, bold=True,
                       color="FFFFFF", bg=header_bg,
                       align='center' if (status_col is not None and j == status_col) else None)
    # body rows
    for row in rows:
        cells = table.add_row().cells
        for j, val in enumerate(row):
            if status_col is not None and j == status_col:
                bg, fg, label = TL.get(str(val).lower(), ("FFFFFF", DARK, val))
                _set_cell_text(cells[j], label, size=body_size, bold=True, color=fg,
                               bg=bg, align='center')
            else:
                _set_cell_text(cells[j], val, size=body_size, color=DARK)
    # column widths
    if widths:
        for j, wdt in enumerate(widths):
            for r in table.rows:
                r.cells[j].width = Twips(int(wdt))
    add_para(doc, "", space_after=4)
    return table


def add_callout(doc, text, bg="EAF7FD", border=BLUE, bold_lead=None, size=10):
    """Single-cell shaded callout box for notices / disclaimers."""
    table = doc.add_table(rows=1, cols=1)
    _set_table_full_width(table)
    cell = table.rows[0].cells[0]
    _set_cell_bg(cell, bg)
    # left accent border
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge, c, s in (('top', border, 4), ('left', border, 18), ('bottom', border, 4), ('right', border, 4)):
        e = OxmlElement('w:' + edge); e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(s)); e.set(qn('w:space'), '0'); e.set(qn('w:color'), c)
        borders.append(e)
    tcPr.append(borders)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    if bold_lead:
        r = p.add_run(bold_lead); _style_run(r, size=size, bold=True, color=DARK)
    r = p.add_run(text); _style_run(r, size=size, color=DARK)
    add_para(doc, "", space_after=4)
    return table


def set_cover_title(doc, title, subtitles):
    """Rewrite the cover banner title and subtitle lines (keeps formatting)."""
    def _replace_text(p, new):
        runs = p.runs
        if not runs:
            r = p.add_run(new); return
        runs[0].text = new
        for r in runs[1:]:
            r.text = ""
    _replace_text(doc.paragraphs[0], title)
    for i, sub in enumerate(subtitles):
        idx = 1 + i
        if idx < len(doc.paragraphs):
            _replace_text(doc.paragraphs[idx], sub)


import zipfile

def _ensure_default_footer(doc):
    """The template body section only defines a 'first-page' footer, so page
    numbers would not show on later pages. Add a default footerReference (to
    footer1) on the body section so the footer/page number appears on every
    body page."""
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is None:
        return
    foot_rid = None
    for rid, rel in doc.part.rels.items():
        if rel.reltype.endswith('/footer') and rel.target_ref.endswith('footer1.xml'):
            foot_rid = rid
            break
    if not foot_rid:
        return
    if any(fr.get(qn('w:type')) == 'default' for fr in sectPr.findall(qn('w:footerReference'))):
        return
    fr = OxmlElement('w:footerReference')
    fr.set(qn('w:type'), 'default'); fr.set(qn('r:id'), foot_rid)
    # place after any existing header/footer references (schema order)
    anchor = None
    for child in sectPr:
        if child.tag in (qn('w:headerReference'), qn('w:footerReference')):
            anchor = child
    if anchor is not None:
        anchor.addnext(fr)
    else:
        sectPr.insert(0, fr)


def finalize(doc, tmp_path, final_path, header_text=None, page_numbers=True):
    """Save doc; optionally patch inner header text and add footer page numbers."""
    if page_numbers:
        _ensure_default_footer(doc)
    doc.save(tmp_path)
    PAGE_PARA = (
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:pPr><w:jc w:val="center"/><w:spacing w:before="40" w:after="0"/></w:pPr>'
        '<w:r><w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>'
        '<w:color w:val="595959"/><w:sz w:val="14"/></w:rPr><w:fldChar w:fldCharType="begin"/></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>'
        '<w:color w:val="595959"/><w:sz w:val="14"/></w:rPr>'
        '<w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>'
        '<w:color w:val="595959"/><w:sz w:val="14"/></w:rPr><w:fldChar w:fldCharType="end"/></w:r></w:p>'
    )
    footers = {'word/footer1.xml'}
    zin = zipfile.ZipFile(tmp_path, 'r')
    zout = zipfile.ZipFile(final_path, 'w', zipfile.ZIP_DEFLATED)
    for item in zin.infolist():
        data = zin.read(item.filename)
        if page_numbers and item.filename in footers:
            data = data.decode('utf-8').replace('</w:ftr>', PAGE_PARA + '</w:ftr>').encode('utf-8')
        if header_text and item.filename == 'word/header1.xml':
            txt = data.decode('utf-8')
            txt = txt.replace('Summary of cover  |  Commercial Combined &amp; Media', header_text)
            data = txt.encode('utf-8')
        zout.writestr(item, data)
    zin.close(); zout.close()
