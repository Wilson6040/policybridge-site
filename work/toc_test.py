"""Test: does LibreOffice populate a TC-field-based TOC (\\f) with page numbers on PDF export?"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK

def add_tc(paragraph, text, level=1):
    p = paragraph._p
    for ic, val in [('begin', None), (None, f' TC "{text}" \\f C \\l "{level}" '), ('end', None)]:
        r = OxmlElement('w:r')
        if ic:
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), ic); r.append(fc)
        else:
            it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = val; r.append(it)
        p.append(r)

def add_toc_field(paragraph):
    p = paragraph._p
    runs = [
        ('begin', None), (None, ' TOC \\f C \\h \\z \\u '), ('separate', None),
        ('text', 'Right-click and choose Update Field to build the contents.'), ('end', None)
    ]
    for kind, val in runs:
        r = OxmlElement('w:r')
        if kind in ('begin', 'separate', 'end'):
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), kind); r.append(fc)
        elif kind == 'text':
            t = OxmlElement('w:t'); t.text = val; r.append(t)
        else:
            it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = val; r.append(it)
        p.append(r)

doc = Document()
h = doc.add_paragraph(); add_toc_field(h)
doc.add_page_break()
p1 = doc.add_paragraph("First Section Heading")
add_tc(p1, "First Section Heading", 1)
doc.add_paragraph("Body " * 400)
doc.add_page_break()
p2 = doc.add_paragraph("Second Section Heading")
add_tc(p2, "Second Section Heading", 1)
doc.add_paragraph("More body " * 400)
doc.add_page_break()
p3 = doc.add_paragraph("Third Section Heading")
add_tc(p3, "Third Section Heading", 1)

# updateFields on open
settings = doc.settings.element
uf = OxmlElement('w:updateFields'); uf.set(qn('w:val'), 'true'); settings.append(uf)
doc.save('/app/work/toc_test.docx')
print("saved toc_test.docx")
