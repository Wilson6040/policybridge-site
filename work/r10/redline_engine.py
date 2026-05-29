# -*- coding: utf-8 -*-
"""
PHASE 01 — TRUE 0223C -> FINAL full tracked-changes redline.

Generates a fresh, holistic Word redline by diffing the ACTUAL paragraph text of
the OLD 0223C baseline against the latest final clean wording. Every addition,
deletion and amendment is emitted as genuine Word tracked changes (w:ins / w:del),
so the document opens in Word/LibreOffice with revisions showing.

Two tracked-change authors are used so they remain separately identifiable:
  * "TMHCC Base Redline (0223C->Final)"  — every ordinary 0223C->final change
  * "UW Review Enhancement"              — the rounds 7-9 enhancements still
                                           awaiting underwriting/legal sign-off
                                           (also yellow-highlighted + annotated).

Outputs:
  /app/outputs/phase-01-full-redline/TMHCC_0223C_to_Final_Full_Redline.docx
  /app/work/r10/redline_stats.json   (stats reused by Phase 02 + QA)
"""
import sys, re, json, difflib
sys.path.insert(0, "/app/work/compare")

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph

import brand2 as B

OLD   = "/app/work/source/OLD_0223C.docx"
FINAL = "/app/deliverables/TMHCC_Media_Combined_0526_FINAL_Clean.docx"
OUT   = "/app/outputs/phase-01-full-redline/TMHCC_0223C_to_Final_Full_Redline.docx"
STATS = "/app/work/r10/redline_stats.json"

AUTHOR_BASE = "TMHCC Base Redline (0223C->Final)"
AUTHOR_UW   = "UW Review Enhancement"
DATE        = "2025-07-01T00:00:00Z"

INK  = RGBColor.from_string(B.INK)
TEAL = RGBColor.from_string(B.TEAL)
GOLD = RGBColor.from_string(B.GOLD)
RUST = RGBColor.from_string(B.RUST)
GREEN = RGBColor.from_string("1E7A46")

# ---- UW enhancement anchors: (substring, enhancement title) --------------
UW_ANCHORS = [
    ("proof of ownership", "Proof of Ownership and Value (claims condition, S1-11)"),
    ("money \u2013 premises", "Money - Premises extension (S8)"),
    ("money \u2013 touring", "Money - Touring, Festivals and Events extension (S8)"),
    ("money waistcoat or money belt", "Money - Touring, Festivals and Events extension (S8)"),
    ("resulting from an assault or violence", "Money - Touring, Festivals and Events extension (S8)"),
    ("discovered within 12 (twelve) months", "Money - Touring, Festivals and Events extension (S8)"),
    ("does not reduce, replace or restrict any cover", "Money extensions (S8)"),
    ("voluntary helper, work-experience or training-scheme", "Employee definition broadened (S13 Legal Expenses)"),
    ("consistent with the meaning of \u201ccomputer system\u201d in section 15", "Computer System definition aligned to S15 (S12)"),
    ("representation costs at investigations", "Representation Costs at Investigations & Inquiries (S12)"),
    ("protection of journalistic sources", "Costs for the Protection of Journalistic Sources (S12)"),
    ("identity of a confidential journalistic source", "Costs for the Protection of Journalistic Sources (S12)"),
    ("criminal and regulatory defence costs", "Criminal & Regulatory Defence Costs cross-reference (S12)"),
    ("distributors and purchasers", "Distributors & Purchasers extension (S12)"),
    ("purchaser, co-producer, licensee or distributor", "Distributors & Purchasers extension (S12)"),
    ("worldwide territory and jurisdiction", "Worldwide Territory & Jurisdiction (Optional) (S12)"),
    ("king\u2019s counsel of the english bar", "King's Counsel determination clause (S12)"),
    ("opinion of such king\u2019s counsel", "King's Counsel determination clause (S12)"),
    ("this write-back shall not extend to any claim for bodily injury", "Pollution negligent-advice write-back (S12)"),
]


def norm(s):
    return re.sub(r"\s+", " ", (s or "")).strip()


def uw_match(text):
    t = norm(text).lower()
    if len(t) < 6:
        return None
    for sub, title in UW_ANCHORS:
        if sub in t:
            return title
    return None


# ---- ordered extraction of (text, style) incl. tables --------------------
def iter_block_paragraphs(doc):
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            tbl = Table(child, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p


def extract(path):
    d = Document(path)
    items = []
    for p in iter_block_paragraphs(d):
        txt = p.text
        if not txt or not txt.strip():
            continue
        style = (p.style.name if p.style else "Normal") or "Normal"
        items.append((txt.strip(), style))
    return items


def style_kind(style):
    s = (style or "").lower()
    if s.startswith("heading 1") or s == "title":
        return "h1"
    if s.startswith("heading 2"):
        return "h2"
    if s.startswith("heading 3") or s == "list heading":
        return "h3"
    if s == "subtitle":
        return "kicker"
    return "body"


# ---- tracked-change run helpers ------------------------------------------
class IdGen:
    def __init__(self):
        self.n = 100
    def next(self):
        self.n += 1
        return self.n

IDS = IdGen()


def _wrap_run(r_el, tag, author):
    parent = r_el.getparent()
    idx = parent.index(r_el)
    w = OxmlElement("w:" + tag)
    w.set(qn("w:id"), str(IDS.next()))
    w.set(qn("w:author"), author)
    w.set(qn("w:date"), DATE)
    parent.insert(idx, w)
    w.append(r_el)
    return w


def _to_deltext(r_el):
    for t in r_el.findall(qn("w:t")):
        t.tag = qn("w:delText")


def _fmt_run(run, kind, sk, uw=False):
    f = run.font
    f.name = B.FONT
    # size / colour by structural kind
    if sk == "h1":
        f.size = Pt(13); f.bold = True; f.color.rgb = TEAL
    elif sk == "h2":
        f.size = Pt(11.5); f.bold = True; f.color.rgb = TEAL
    elif sk == "h3":
        f.size = Pt(10.5); f.bold = True; f.color.rgb = GOLD
    elif sk == "kicker":
        f.size = Pt(11); f.bold = True; f.color.rgb = GOLD
    else:
        f.size = Pt(9.5); f.color.rgb = INK
    # revision visual cues (Word also recolours by author)
    if kind == "del":
        f.strike = True
        f.color.rgb = RUST
    elif kind == "ins":
        f.underline = True
        if not uw and sk == "body":
            f.color.rgb = GREEN
    if uw:
        f.highlight_color = WD_COLOR_INDEX.YELLOW


def emit_paragraph(doc, text, kind, sk, author=AUTHOR_BASE, uw=False, annotate=False):
    """kind in {equal, ins, del}."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if sk in ("h1",):
        pf.space_before = Pt(10); pf.space_after = Pt(3); pf.keep_with_next = True
    elif sk in ("h2", "h3", "kicker"):
        pf.space_before = Pt(6); pf.space_after = Pt(2); pf.keep_with_next = True
    else:
        pf.space_after = Pt(3)
    run = p.add_run(text)
    _fmt_run(run, kind, sk, uw=uw)
    if kind == "ins":
        _wrap_run(run._r, "ins", author)
    elif kind == "del":
        _to_deltext(run._r)
        _wrap_run(run._r, "del", author)
    if annotate and uw:
        a = p.add_run("  [UW REVIEW ENHANCEMENT \u2014 awaiting underwriting/legal sign-off; not accepted]")
        a.font.name = B.FONT; a.font.size = Pt(7.5); a.font.italic = True
        a.font.color.rgb = RUST
    return p


def emit_inline_diff(doc, old_text, new_text, sk):
    """Word-level inline redline for an amended paragraph (one old -> one new)."""
    ow = re.findall(r"\S+\s*", old_text)
    nw = re.findall(r"\S+\s*", new_text)
    sm = difflib.SequenceMatcher(a=[w.strip().lower() for w in ow],
                                 b=[w.strip().lower() for w in nw], autojunk=False)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            run = p.add_run("".join(ow[i1:i2])); _fmt_run(run, "equal", sk)
        elif tag == "delete":
            run = p.add_run("".join(ow[i1:i2])); _fmt_run(run, "del", sk)
            _to_deltext(run._r); _wrap_run(run._r, "del", AUTHOR_BASE)
        elif tag == "insert":
            run = p.add_run("".join(nw[j1:j2])); _fmt_run(run, "ins", sk)
            _wrap_run(run._r, "ins", AUTHOR_BASE)
        elif tag == "replace":
            run = p.add_run("".join(ow[i1:i2])); _fmt_run(run, "del", sk)
            _to_deltext(run._r); _wrap_run(run._r, "del", AUTHOR_BASE)
            run = p.add_run("".join(nw[j1:j2])); _fmt_run(run, "ins", sk)
            _wrap_run(run._r, "ins", AUTHOR_BASE)
    return p


# ==========================================================================
def main():
    old_items = extract(OLD)
    new_items = extract(FINAL)
    old_keys = [norm(t).lower() for t, _ in old_items]
    new_keys = [norm(t).lower() for t, _ in new_items]

    sm = difflib.SequenceMatcher(a=old_keys, b=new_keys, autojunk=False)
    opcodes = sm.get_opcodes()

    doc = Document()
    B.set_normal_style(doc)
    B.setup_page(doc, landscape=False, margin_cm=1.8)
    B.add_logo_header(doc, left_text="TMHCC Media & Music Combined \u2014 Full Redline 0223C \u2192 Final")
    B.add_footer_pagenum(doc, note="FULL TRACKED-CHANGES REDLINE: TMHCC 0223C \u2192 final wording (commit 58b8340). "
                                   "Open with Track Changes shown. Base changes vs UW Review Enhancements are separate authors. "
                                   "Subject to TMHCC legal/underwriting sign-off.")

    # ---- intro / how-to-read page ----
    B.h1(doc, "Full Redline \u2014 TMHCC 0223C \u2192 Final Wording")
    B.para(doc, "This document is a complete, holistic tracked-changes comparison generated afresh from the "
                "TMHCC 0223C baseline wording against the latest final wording (commit 58b8340). It is NOT a "
                "summary of recent enhancements \u2014 it shows every textual difference between the two wordings.",
           size=10)
    B.callout(doc, "How to read this redline:",
              "Open in Microsoft Word with Review \u25b8 All Markup. Deleted 0223C text shows struck-through (red); "
              "inserted final-wording text shows underlined. Two tracked-change authors are used so they remain "
              "separately identifiable: \u201cTMHCC Base Redline (0223C\u2192Final)\u201d for ordinary changes, and "
              "\u201cUW Review Enhancement\u201d (yellow-highlighted) for the rounds 7\u20139 enhancements that are still "
              "awaiting underwriting/legal sign-off and must NOT be accepted without sign-off. A separate Enhancement "
              "Register lists every UW item.")
    B.rich(doc, [("Legend:   ", {"bold": True, "color": B.TEAL}),
                 ("deleted from 0223C", {"color": B.RUST}), ("      ", {}),
                 ("inserted into final wording", {"color": "1E7A46"}), ("      ", {}),
                 ("UW Review Enhancement (yellow highlight)", {"color": B.GOLD})], size=9)
    doc.add_paragraph().add_run().add_break()

    # ---- emit redline ----
    stats = dict(equal=0, insert=0, delete=0, replace=0, ins_paras=0, del_paras=0,
                 inline_amend=0, uw_paras=0)
    uw_hits = {}

    def record_uw(title):
        uw_hits[title] = uw_hits.get(title, 0) + 1
        stats["uw_paras"] += 1

    for tag, i1, i2, j1, j2 in opcodes:
        if tag == "equal":
            stats["equal"] += (i2 - i1)
            for k in range(j1, j2):
                txt, style = new_items[k]
                emit_paragraph(doc, txt, "equal", style_kind(style))
        elif tag == "insert":
            stats["insert"] += (j2 - j1)
            uw_active = None
            for k in range(j1, j2):
                txt, style = new_items[k]
                sk = style_kind(style)
                title = uw_match(txt)
                if title:
                    uw_active = title
                    emit_paragraph(doc, txt, "ins", sk, author=AUTHOR_UW, uw=True,
                                   annotate=True)
                    record_uw(title)
                elif sk in ("h1", "h2", "h3", "kicker"):
                    uw_active = None
                    emit_paragraph(doc, txt, "ins", sk, author=AUTHOR_BASE)
                    stats["ins_paras"] += 1
                elif uw_active:
                    emit_paragraph(doc, txt, "ins", sk, author=AUTHOR_UW, uw=True)
                    record_uw(uw_active)
                else:
                    emit_paragraph(doc, txt, "ins", sk, author=AUTHOR_BASE)
                    stats["ins_paras"] += 1
        elif tag == "delete":
            stats["delete"] += (i2 - i1)
            for k in range(i1, i2):
                txt, style = old_items[k]
                emit_paragraph(doc, txt, "del", style_kind(style), author=AUTHOR_BASE)
                stats["del_paras"] += 1
        elif tag == "replace":
            stats["replace"] += 1
            # one-to-one short amendments -> inline word diff; else block del + block ins
            if (i2 - i1) == 1 and (j2 - j1) == 1:
                ot, ostyle = old_items[i1]
                nt, nstyle = new_items[j1]
                title = uw_match(nt)
                if title:
                    emit_paragraph(doc, nt, "ins", style_kind(nstyle), author=AUTHOR_UW,
                                   uw=True, annotate=True)
                    emit_paragraph(doc, ot, "del", style_kind(ostyle), author=AUTHOR_BASE)
                    record_uw(title)
                else:
                    emit_inline_diff(doc, ot, nt, style_kind(nstyle))
                    stats["inline_amend"] += 1
            else:
                for k in range(i1, i2):
                    txt, style = old_items[k]
                    emit_paragraph(doc, txt, "del", style_kind(style), author=AUTHOR_BASE)
                    stats["del_paras"] += 1
                uw_active = None
                for k in range(j1, j2):
                    txt, style = new_items[k]
                    sk = style_kind(style)
                    title = uw_match(txt)
                    if title:
                        uw_active = title
                        emit_paragraph(doc, txt, "ins", sk, author=AUTHOR_UW, uw=True, annotate=True)
                        record_uw(title)
                    elif sk in ("h1", "h2", "h3", "kicker"):
                        uw_active = None
                        emit_paragraph(doc, txt, "ins", sk, author=AUTHOR_BASE)
                        stats["ins_paras"] += 1
                    elif uw_active:
                        emit_paragraph(doc, txt, "ins", sk, author=AUTHOR_UW, uw=True)
                        record_uw(uw_active)
                    else:
                        emit_paragraph(doc, txt, "ins", sk, author=AUTHOR_BASE)
                        stats["ins_paras"] += 1

    # ---- section-level structural change detection (Heading 1) ----
    old_h1 = [norm(t) for t, s in old_items if style_kind(s) == "h1"]
    new_h1 = [norm(t) for t, s in new_items if style_kind(s) == "h1"]
    old_set = {h.lower(): h for h in old_h1}
    new_set = {h.lower(): h for h in new_h1}
    added_sections = [new_set[k] for k in new_set if k not in old_set]
    removed_sections = [old_set[k] for k in old_set if k not in new_set]

    stats["old_paras"] = len(old_items)
    stats["new_paras"] = len(new_items)
    stats["old_h1"] = len(old_h1)
    stats["new_h1"] = len(new_h1)
    stats["added_section_headings"] = added_sections
    stats["removed_section_headings"] = removed_sections
    stats["uw_enhancements_detected"] = uw_hits

    B.save_doc(doc, OUT)
    with open(STATS, "w") as f:
        json.dump(stats, f, indent=2)

    print("SAVED:", OUT)
    print("STATS:", json.dumps({k: v for k, v in stats.items()
                                 if k not in ("added_section_headings", "removed_section_headings")}, indent=2))
    print("UW enhancements detected:", json.dumps(uw_hits, indent=2))
    print("Added H1 headings (sample):", added_sections[:25])
    print("Removed H1 headings (sample):", removed_sections[:25])


if __name__ == "__main__":
    main()
