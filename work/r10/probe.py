"""Probe OLD_0223C.docx vs final clean docx for the redline engine."""
import sys
from docx import Document

OLD = "/app/work/source/OLD_0223C.docx"
FINAL = "/app/deliverables/TMHCC_Media_Combined_0526_FINAL_Clean.docx"


def iter_block_paras(doc):
    """Return list of (text, style) for all paragraphs incl. inside tables, in order."""
    out = []
    for p in doc.paragraphs:
        out.append((p.text, p.style.name if p.style else ""))
    # tables
    tcount = 0
    for t in doc.tables:
        tcount += 1
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        out.append((p.text, "TableCell"))
    return out, tcount


for label, path in [("OLD_0223C", OLD), ("FINAL_CLEAN", FINAL)]:
    d = Document(path)
    paras, tcount = iter_block_paras(d)
    nonempty = [p for p in paras if p[0].strip()]
    print(f"=== {label} : {path}")
    print(f"  total body paragraphs: {len(d.paragraphs)} | tables: {tcount} | nonempty (incl table): {len(nonempty)}")
    # style histogram
    from collections import Counter
    styles = Counter(s for _, s in paras)
    print("  styles:", dict(styles.most_common(12)))
    # sample headings (Heading styles)
    heads = [t for t, s in paras if s and s.startswith("Heading") and t.strip()][:25]
    print("  sample headings:")
    for h in heads[:25]:
        print("     -", h[:90])
    print()
