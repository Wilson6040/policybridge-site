# -*- coding: utf-8 -*-
"""
Round 8 (EMERGENT DISPATCH v3) — TMHCC wording tracked-change builder.

Produces:
  baseline_clean.docx          (accept any existing tracked changes -> clean baseline)
  ..._TrackedChanges.docx      (baseline + NEW discrete tracked changes; each amendment self-contained)
  ..._Clean.docx               (accept-all of the new tracked changes)

Every amendment is a discrete, individually accept/reject-able tracked change.
Definition replacements show OLD text deleted and NEW text inserted (old visible).
"""
import copy
from docx import Document
from lxml import etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
def qn(t): return f"{{{W}}}{t}"
def ln(el): return etree.QName(el).localname

AUTHOR = "TMHCC Wording Review (v3)"
DATE = "2025-07-01T00:00:00Z"
_id = [20000]
def nid():
    _id[0] += 1
    return str(_id[0])

SRC = "/app/work/r8/source_uploaded.docx"
OUT_TRACK = "/app/work/r8/TMHCC_Media_Combined_0526_FINAL_TrackedChanges.docx"
OUT_CLEAN = "/app/work/r8/TMHCC_Media_Combined_0526_FINAL_Clean.docx"
OUT_BASE  = "/app/work/r8/baseline_clean.docx"

# ---------- low-level run builders ----------
def new_run(text, bold=False, rpr_src=None):
    r = etree.Element(qn("r"))
    rpr = None
    if rpr_src is not None:
        src = rpr_src.find(qn("rPr"))
        if src is not None:
            rpr = copy.deepcopy(src)
            # strip any existing ins/del marks inside copied rPr
            for bad in rpr.findall(qn("ins")) + rpr.findall(qn("del")):
                rpr.remove(bad)
    if bold:
        if rpr is None:
            rpr = etree.Element(qn("rPr"))
        if rpr.find(qn("b")) is None:
            rpr.append(etree.Element(qn("b")))
    if rpr is not None:
        r.append(rpr)
    t = etree.SubElement(r, qn("t"))
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return r

def wrap_ins(children):
    ins = etree.Element(qn("ins"))
    ins.set(qn("id"), nid()); ins.set(qn("author"), AUTHOR); ins.set(qn("date"), DATE)
    for c in children:
        ins.append(c)
    return ins

def wrap_del(runs):
    delel = etree.Element(qn("del"))
    delel.set(qn("id"), nid()); delel.set(qn("author"), AUTHOR); delel.set(qn("date"), DATE)
    for r in runs:
        for t in r.findall(qn("t")):
            t.tag = qn("delText")
        delel.append(r)
    return delel

def mark_para_inserted(p):
    pPr = p.find(qn("pPr"))
    if pPr is None:
        pPr = etree.Element(qn("pPr")); p.insert(0, pPr)
    rPr = pPr.find(qn("rPr"))
    if rPr is None:
        rPr = etree.SubElement(pPr, qn("rPr"))
    ins = etree.SubElement(rPr, qn("ins"))
    ins.set(qn("id"), nid()); ins.set(qn("author"), AUTHOR); ins.set(qn("date"), DATE)

def new_para(style, run_segments, rpr_src=None):
    """run_segments: list of (text, bold). Returns w:p with inserted runs + inserted para mark."""
    p = etree.Element(qn("p"))
    pPr = etree.SubElement(p, qn("pPr"))
    ps = etree.SubElement(pPr, qn("pStyle")); ps.set(qn("val"), style)
    runs = [new_run(txt, bold, rpr_src=rpr_src) for (txt, bold) in run_segments]
    p.append(wrap_ins(runs))
    mark_para_inserted(p)
    return p

# ---------- accept helpers ----------
def accept_all(body):
    for ins in list(body.iter(qn("ins"))):
        parent = ins.getparent()
        if parent is None:
            continue
        if ln(parent) == "rPr":
            parent.remove(ins)
        else:
            idx = list(parent).index(ins)
            for c in reversed(list(ins)):
                parent.insert(idx, c)
            parent.remove(ins)
    for d in list(body.iter(qn("del"))):
        if d.getparent() is not None:
            d.getparent().remove(d)
    for tag in ("rPrChange", "pPrChange"):
        for e in list(body.iter(qn(tag))):
            if e.getparent() is not None:
                e.getparent().remove(e)

# ---------- find helpers (iterate ALL w:p incl. tables) ----------
def all_p(doc):
    return list(doc.element.body.iter(qn("p")))

def ptext(p):
    return "".join((t.text or "") for t in p.iter(qn("t")))

def find_p(doc, substr):
    hits = [p for p in all_p(doc) if substr in ptext(p)]
    assert len(hits) == 1, f"anchor not unique ({len(hits)}): {substr!r}"
    return hits[0]

def runs_of(p):
    return [c for c in p if ln(c) == "r"]

def insert_blocks_after(anchor_p, blocks):
    """blocks: list of (style, segments). Insert as new inserted paragraphs after anchor, in order."""
    cur = anchor_p
    rpr_src = None
    # try to inherit run formatting from anchor's first run
    rs = runs_of(anchor_p)
    if rs:
        rpr_src = rs[0]
    for (style, segs) in blocks:
        p = new_para(style, segs, rpr_src=rpr_src)
        cur.addnext(p)
        cur = p
    return cur

def replace_para_text(p, new_segments):
    """Delete all existing runs of p (as tracked deletion) and insert new runs (tracked insertion)."""
    existing = runs_of(p)
    rpr_src = existing[0] if existing else None
    delel = wrap_del(existing)  # detaches runs into del, converts t->delText
    insel = wrap_ins([new_run(t, b, rpr_src=rpr_src) for (t, b) in new_segments])
    p.append(delel)
    p.append(insel)

# ====================================================================
# STEP 1 — baseline (accept any existing tracked changes)
# ====================================================================
doc = Document(SRC)
accept_all(doc.element.body)
doc.save(OUT_BASE)
print("baseline_clean.docx saved")

# ====================================================================
# STEP 2 — apply NEW discrete tracked changes
# ====================================================================
doc = Document(OUT_BASE)
applied = []

# ---- A4 (round 9): TWO Schedule-referenced money extensions (Premises + Touring/Festivals/Events),
#       amalgamating the Entertainment Elite event-money section + the 0523 "Venue" enhancement.
#       Preserves all existing Money covers/exclusions (added after Loss of Keys body). ----
a = find_p(doc, "the cost of replacement keys or lock mechanisms")
insert_blocks_after(a, [
 ("Heading3", [("Money \u2013 Premises", False)]),
 ("BodyText", [
   ("This Sub-Section is extended to cover Damage to ", False), ("Money", True),
   (" belonging to the ", False), ("Insured", True),
   (", or for which the ", False), ("Insured", True),
   (" is responsible, whilst at the Premises, up to the sub-limit shown against this Extension in the Schedule. This Extension does not reduce, replace or restrict any cover otherwise provided by this Sub-Section and is subject otherwise to the Definitions, Conditions and Exclusions of this Sub-Section.", False),
 ]),
 ("Heading3", [("Money \u2013 Touring, Festivals and Events", False)]),
 ("BodyText", [
   ("This Sub-Section is extended to cover Damage to ", False), ("Money", True),
   (" belonging to the ", False), ("Insured", True),
   (", or for which the ", False), ("Insured", True),
   (" is responsible, whilst in transit, in temporary storage, in a locked safe, cash point or container, in the personal custody of the ", False),
   ("Insured", True),
   (" or an authorised ", False), ("Employee", True),
   (", or at any Venue, temporary venue, festival site, box office or other location used in connection with touring, live events, performances, productions, exhibitions, conferences, product launches or other music or entertainment activities undertaken by the ", False),
   ("Insured", True),
   (", up to the sub-limit shown against this Extension in the Schedule. This Extension includes:", False),
 ]),
 ("BodyText", [
   ("(a) Damage to any safe, strongroom, cash point, bag, container, money waistcoat or money belt used to carry or hold such ", False),
   ("Money", True), (";", False),
 ]),
 ("BodyText", [
   ("(b) Damage to clothing, personal effects and personal ", False), ("Money", True),
   (" belonging to the ", False), ("Insured", True),
   (" or any partner, director or ", False), ("Employee", True),
   (" resulting from an assault or violence, or the threat of assault or violence, in an attempt to steal ", False),
   ("Money", True), ("; and", False),
 ]),
 ("BodyText", [
   ("(c) loss of ", False), ("Money", True),
   (" belonging to the ", False), ("Insured", True),
   (", or for which the ", False), ("Insured", True),
   (" is responsible, occasioned by the dishonesty of any ", False), ("Employee", True),
   (" of the ", False), ("Insured", True),
   (", committed during the Period of Insurance and discovered within 12 (twelve) months of the date of the theft, up to the sub-limit shown against this Extension in the Schedule.", False),
 ]),
 ("BodyText", [
   ("This Extension does not reduce, replace or restrict any cover otherwise provided by this Sub-Section and is subject otherwise to the Definitions, Conditions and Exclusions of this Sub-Section.", False),
 ]),
])
applied.append("A4 Money RESTRUCTURED into two Schedule-referenced extensions: 'Money \u2013 Premises' and "
               "'Money \u2013 Touring, Festivals and Events' (amalgamates Entertainment Elite event-money + 0523 Venue)")

# ---- A5: broaden S13 Employee definition (replace narrow body) ----
a = find_p(doc, "alleges they have entered into a contract of service")
replace_para_text(a, [
  ("Any natural person who is, or who alleges they are or were, working for ", False), ("You", True),
  (" under ", False), ("Your", True),
  (" direction, control or supervision in connection with the ", False), ("Business", True),
  (", including: (a) any person under a contract of service or apprenticeship with ", False), ("You", True),
  ("; (b) any person employed on a labour-only basis, or self-employed and supplying labour only under a contract for services; (c) any person hired or borrowed by ", False),
  ("You", True),
  (" from another employer; and (d) any voluntary helper, work-experience or training-scheme participant, secondee, student or prospective employee being assessed as to their suitability for employment.", False),
])
applied.append("A5 Employee definition broadened (S13 Legal Expenses) [old->new visible]")

# ---- A6: align S12 Computer System definition to the broader S15/general definition (FLAG S12 sign-off) ----
a = find_p(doc, "Any computer, data processing equipment, media or part thereof")
replace_para_text(a, [
  ("Any computer or computing, electronic, wireless, web or similar system (including any associated input, output, data storage, networking or back-up device or facility), and all computer hardware, software, firmware, microcode, data, communications systems, networks, protocols, electronic documents, websites, intranet, extranet and cloud or off-line storage facilities, used to process, store, transmit or retrieve data or information in any analogue, digital, electronic or wireless format, owned, operated, leased, licensed or utilised by the ", False),
  ("Insured", True),
  (" or any other party. This definition is intended to be consistent with the meaning of \u201cComputer System\u201d in Section 15 (CyberGuard\u2122 (Cyber Liability)).", False),
])
applied.append("A6 S12 Computer System definition aligned to S15 [SIGN-OFF] [old->new visible]")

# ---- B2 + B3 + B4: new S12 additional covers (after IP-pursuit clause) (FLAG sign-off) ----
a = find_p(doc, "a measurable loss and a reasonable prospect of success")
insert_blocks_after(a, [
 ("Heading3", [("Representation Costs at Investigations and Inquiries", False)]),
 ("BodyText", [
   ("Up to a maximum of GBP 25,000 (twenty-five thousand pounds sterling) in the aggregate in the Period of Insurance, for the reasonable costs and expenses incurred by the ", False),
   ("Insured", True),
   (", with the ", False), ("Insurer\u2019s", True),
   (" prior written consent, in respect of attendance at or representation before any official examination, investigation, inquiry or other proceeding arising from the ", False),
   ("Insured\u2019s", True),
   (" Media Business Services and which is likely to give rise to a claim under Insuring Clause 1. (Indemnity). This cover forms part of and does not increase the Indemnity Limit.", False),
 ]),
 ("Heading3", [("Costs for the Protection of Journalistic Sources", False)]),
 ("BodyText", [
   ("Up to the sub-limit stated in the Schedule against this Extension, for the reasonable legal costs and expenses incurred by the ", False),
   ("Insured", True),
   (", with the ", False), ("Insurer\u2019s", True),
   (" prior written consent, in opposing or responding to any subpoena, court order or other legal compulsion requiring the ", False),
   ("Insured", True),
   (" to disclose the identity of a confidential journalistic source, where such disclosure arises from the ", False),
   ("Insured\u2019s", True),
   (" Media Business Services. The ", False), ("Insurer", True),
   (" will only provide such cover where, in the opinion of counsel, there are reasonable grounds to oppose disclosure. This cover forms part of and does not increase the Indemnity Limit.", False),
 ]),
 ("Heading3", [("Criminal and Regulatory Defence Costs", False)]),
 ("BodyText", [
   ("In addition to the Data Protection Defence Costs provided above, the ", False), ("Insured", True),
   (" is reminded that defence of criminal and regulatory proceedings may also be available under Section 13 (Commercial Legal Expenses) and Section 14 (Management Liability), subject to the terms of those Sections. The ", False),
   ("Insurer", True),
   (" will further contribute, up to the sub-limit stated in the Schedule against this Extension and with the ", False),
   ("Insurer\u2019s", True),
   (" prior written consent, towards the reasonable defence costs of any criminal or regulatory proceeding brought against the ", False),
   ("Insured", True),
   (" arising directly from the conduct of the ", False), ("Insured\u2019s", True),
   (" Media Business Services, save to the extent that cover is provided under Section 13 or Section 14. This cover forms part of and does not increase the Indemnity Limit and excludes fines, penalties and any costs incurred after a plea or finding of guilt.", False),
 ]),
])
applied.append("B2 Representation Costs (S12) [SIGN-OFF] GBP25k (Tysers 8.15)")
applied.append("B3 Journalistic Source-Protection Costs (S12) [SIGN-OFF] (Tysers 8.13)")
applied.append("B4 Criminal & Regulatory Defence Costs cross-reference + contribution (S12) [SIGN-OFF]")

# ---- B1 + B5: new S12 Extensions (after Indemnity to Principals body) (FLAG sign-off) ----
a = find_p(doc, "indemnify any Principal with whom the Insured has entered")
insert_blocks_after(a, [
 ("Heading3", [("Distributors and Purchasers", False)]),
 ("BodyText", [
   ("The indemnity provided by Insuring Clause 1. (Indemnity) is extended to apply to any purchaser, co-producer, licensee or distributor of Media Material produced or supplied by the ", False),
   ("Insured", True),
   (", but only in respect of claims arising from such Media Material and only where the ", False), ("Insured", True),
   (" is contractually obliged to provide such indemnity. The ", False), ("Insurer\u2019s", True),
   (" total liability under this Extension shall not exceed five (5) times the Indemnity Limit in the aggregate during the Period of Insurance for all claims arising from each production of Media Material so purchased, co-produced, licensed or distributed. Cover under this Extension is conditional upon the ", False),
   ("Insurer", True),
   (" retaining the right to conduct the defence and settlement of any such claim, and upon any claim being first made against, and notified to the ", False),
   ("Insurer", True),
   (" by, the ", False), ("Insured", True),
   (" during the Period of Insurance.", False),
 ]),
 ("Heading3", [("Worldwide Territory and Jurisdiction (Optional)", False)]),
 ("BodyText", [
   ("This Optional Extension applies only if stated as operative in the Schedule. Where stated as operative, the Geographical Limits applicable to this Section are amended to anywhere in the world and the Jurisdiction applicable to this Section is amended to worldwide, and the exclusions of claims brought, or in which the governing law is contended to be, outside the Jurisdiction shall not apply. Cover in respect of claims brought in, or determined under the laws of, the United States of America or Canada (or any territory or protectorate thereof) applies only if \u201cUSA/Canada Jurisdiction\u201d is additionally stated as operative in the Schedule.", False),
 ]),
])
applied.append("B1 Distributors & Purchasers Extension (S12) [SIGN-OFF] 5x limit (Tysers 8.7)")
applied.append("B5 Worldwide Territory & Jurisdiction (Optional) Extension (S12) [SIGN-OFF] (was Gap 2)")

# ---- B6: Senior/King's Counsel determination clause (after the S12 dispute/mediation clause) ----
a = find_p(doc, "In respect of Section 12 only any dispute concerning the interpretation")
insert_blocks_after(a, [
 ("BodyText", [
   ("In respect of Section 12 (Media Liability) only, if the ", False), ("Insured", True),
   (" and the ", False), ("Insurer", True),
   (" disagree as to whether or how any claim should be defended, settled or pursued, either party may require the matter to be referred to a King\u2019s Counsel of the English Bar to be mutually agreed between the ", False),
   ("Insured", True), (" and the ", False), ("Insurer", True),
   (" (or, failing agreement within 14 (fourteen) working days, to be appointed by the Chairman of the Bar Council or their representative). The opinion of such King\u2019s Counsel as to the conduct of the matter shall be binding on both parties. In giving such opinion the King\u2019s Counsel shall have regard to the interests of both the ", False),
   ("Insured", True), (" and the ", False), ("Insurer", True),
   (", and the costs of the reference shall be allocated by the King\u2019s Counsel on a fair and equitable basis.", False),
 ]),
])
applied.append("B6 King's Counsel determination clause (S12 dispute) [SIGN-OFF] (Yutree)")

# ---- B7: Pollution negligent-advice write-back (after S12 Seepage & Pollution exclusion) ----
a = find_p(doc, "in any way involving seepage, Pollution or contamination of any kind")
insert_blocks_after(a, [
 ("BodyText", [
   ("provided that this Exclusion shall not apply to any claim arising solely from a negligent act, negligent error or negligent omission in the provision of the ", False),
   ("Insured\u2019s", True),
   (" Media Business Services (and not from any seepage, Pollution or contamination caused by the ", False),
   ("Insured\u2019s", True),
   (" own operations, premises or activities), and provided further that this write-back shall not extend to any claim for Bodily Injury, property damage, or the costs of removing, nullifying, cleaning up or remediating any pollutant. The ", False),
   ("Insurer\u2019s", True),
   (" liability under this write-back shall not exceed the sub-limit stated in the Schedule against this write-back.", False),
 ]),
])
applied.append("B7 Pollution negligent-advice write-back (S12) [SIGN-OFF] (Yutree)")

doc.save(OUT_TRACK)
print("TrackedChanges.docx saved")
print("Applied %d discrete amendments:" % len(applied))
for x in applied:
    print("  -", x)

# ====================================================================
# STEP 3 — clean final (accept all new tracked changes)
# ====================================================================
doc2 = Document(OUT_TRACK)
accept_all(doc2.element.body)
doc2.save(OUT_CLEAN)
print("Clean.docx saved")
