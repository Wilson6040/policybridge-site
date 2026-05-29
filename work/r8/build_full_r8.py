# -*- coding: utf-8 -*-
"""D6 — CLIENT / BROKER-FACING Full Coverage Comparison.
A selection/sales tool: TMHCC strengths and where TMHCC matches or beats the market.
Contains NO gap analysis, NO internal deficiencies, NO sub-limit working, NO sign-off notes.
"""
import sys
sys.path.insert(0, '/app/work/compare')
sys.path.insert(0, '/app/work/r8')
from docx import Document
from docx.shared import Pt
import brand2 as b
import cmpdata as d

W_HEADERS = ["TMHCC", "Tysers\n(Zurich)", "Yutree\n(AXA)", "Liberty", "Allianz", "AXA XL"]


def comment_cell(cell, text):
    cell.text = ""
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text); b._set_run(r, size=8, color=b.INK)


def section_matrix(doc, rows):
    headers = ["Section of cover"] + W_HEADERS + ["What this means for you"]
    widths = [5.0] + [1.7] * 6 + [10.8]
    t = b.make_table(doc, headers, widths)
    for label, statuses, verdict, comment in rows:
        cells = t.add_row().cells
        b.text_cell(cells[0], label, size=8.3, bold=True, color=b.TEAL)
        for j, key in enumerate(statuses):
            b.status_cell(cells[1 + j], key)
        comment_cell(cells[7], comment)
    b.zebra(t); b.spacer(doc, 6)
    return t


def build():
    doc = Document()
    b.set_normal_style(doc)
    b.setup_page(doc, landscape=True, margin_cm=1.5)
    b.add_logo_header(doc)
    b.add_footer_pagenum(doc, note="Client & broker reference. Indicative comparison of policy wordings — always refer to the full policy wordings and Schedule. Cover is subject to the policy terms, conditions and limits.")

    b.cover_page(
        doc,
        doc_kicker="Coverage Comparison — Media & Entertainment",
        title="Why TMHCC — Coverage Comparison",
        subtitle="How the Tokio Marine HCC Media & Music Combined wording compares, section by section, with five other entertainment-market wordings (Tysers/Zurich, Yutree/AXA, Liberty, Allianz and AXA XL).",
        meta_lines=[
            "A client- and broker-facing guide to the breadth of the TMHCC wording. Indicative comparison for selection purposes only — always refer to the full policy wordings and the Schedule.",
            "Prepared from the wordings supplied; competitor capacity insurers noted where stated.",
        ],
        wordings=[f"{w['name']} — {w['full']} ({w['insurer']}); {w['sections']}" for w in d.WORDINGS],
    )

    # 1. Why TMHCC
    doc.add_page_break()
    b.h1(doc, "Why TMHCC — in short", num="1")
    b.para(doc, "The TMHCC Media & Music Combined wording is the broadest in this peer group. It is the only wording that brings all fifteen covers together in a single contract, and the only one to include Loss of Licence, Commercial Legal Expenses, Management Liability and a genuine standalone Cyber-liability section. Three of the five competitors (Liberty, Allianz and AXA XL) are property / business-interruption / liability packages with no media, professional-indemnity, production-indemnity, cyber, legal-expenses or management-liability cover at all.", align='just')
    b.bullet(doc, "All fifteen covers in one place — fewer gaps, fewer separate policies to manage.", bold_lead="One-stop wording:  ")
    b.bullet(doc, "Loss of Licence, Commercial Legal Expenses, Management Liability and standalone CyberGuard™ (Cyber Liability) — not offered by any of the five competitors reviewed.", bold_lead="Four covers no competitor offers:  ")
    b.bullet(doc, "A broad Media Liability / Professional Indemnity section with reputation management, withdrawal-of-content, data-protection defence costs, virus cover and IP-pursuit costs — plus recently added market-matching enhancements.", bold_lead="Strong media cover:  ")
    b.bullet(doc, "A clean exclusion structure that keeps the cyber/date exclusion away from the liability sections, avoiding the ‘silent-cyber strip’ seen elsewhere.", bold_lead="Clear protection:  ")

    # 2. Wordings compared
    b.h1(doc, "The wordings compared", num="2")
    t = b.make_table(doc, ["Wording", "Full title", "Capacity / insurer", "Structure"], [3.6, 9.5, 7.0, 5.9])
    for w in d.WORDINGS:
        cells = t.add_row().cells
        b.text_cell(cells[0], w['name'], size=8.5, bold=True, color=b.TEAL)
        b.text_cell(cells[1], w['full'], size=8.3)
        b.text_cell(cells[2], w['insurer'], size=8.3)
        b.text_cell(cells[3], w['sections'], size=8.3)
    b.zebra(t); b.spacer(doc, 4)
    b.callout(doc, "How to read the matrix:", "✓ = cover clearly present · ◐ = partial / conditional / sub-limited · ✗ = not covered / not identified · ? = unclear. The right-hand column explains, in plain language, what each row means for you.", bg="FBF1DA", border=b.GOLD, lead_color="8A6A1E")

    # 3. At a glance
    b.h1(doc, "Cover at a glance", num="3")
    b.para(doc, "Section-by-section availability across all six wordings.")
    b.legend_bar(doc)
    section_matrix(doc, d.SECTION_ROWS)

    # 4. TMHCC strengths
    doc.add_page_break()
    b.h1(doc, "Where TMHCC stands out", num="4")
    b.para(doc, "The areas where the TMHCC wording is demonstrably stronger, broader or clearer than the competitor wordings reviewed:")
    t = b.make_table(doc, ["Strength", "Why it matters to you"], [7.0, 19.0])
    for title, why in d.TMHCC_STRENGTHS:
        cells = t.add_row().cells
        b.text_cell(cells[0], title, size=8.6, bold=True, color=b.TEAL)
        b.text_cell(cells[1], why, size=8.5)
    b.zebra(t); b.spacer(doc, 6)

    # 5. Recent enhancements (positive framing, options)
    b.h1(doc, "Recent enhancements to the Media Liability section", num="5")
    b.para(doc, "The Media Liability section has been further strengthened so that it stands with the strongest media wordings in the market. The following enhancements are available (certain options apply where shown in the Schedule):")
    for lead, txt in [
        ("Distributors & Purchasers:", "indemnity can extend to purchasers, co-producers, licensees and distributors of your media material."),
        ("Worldwide / USA-Canada option:", "an optional worldwide territory (with USA/Canada separately selectable) where shown in the Schedule."),
        ("Representation costs:", "costs of attending official investigations and inquiries arising from your media activities."),
        ("Source-protection costs:", "legal costs of resisting an order to disclose a confidential journalistic source."),
        ("Criminal & regulatory defence:", "a clear route to criminal/regulatory defence via Sections 13 and 14, plus a contribution within the media section."),
        ("King’s Counsel determination:", "a fair, binding way to resolve any defence/settlement disagreement."),
    ]:
        b.bullet(doc, txt, bold_lead=lead)
    b.para(doc, "Touring & entertainment money cover and a broadened definition of ‘Employee’ in the Legal Expenses section have also been added.", size=8.6, italic=True, color=b.GREY)

    # 6. Competitor profiles
    b.h1(doc, "The competitor wordings at a glance", num="6")
    for key in d.COMPETITOR_KEYS:
        prof = d.COMPETITOR_PROFILES[key]
        b.h2(doc, prof['title'])
        b.para(doc, prof['shape'], italic=True, color=b.GREY, size=9)
        for s in prof['strengths']:
            b.bullet(doc, s, bold_lead="Strength:  ", color=b.INK)
        for g in prof['gaps']:
            b.bullet(doc, g, bold_lead="Not in this wording:  ", color=b.INK)
        b.spacer(doc, 4)

    # 7. Summary
    b.h1(doc, "In summary", num="7")
    b.callout(doc, "The bottom line:", "For a media or entertainment client that wants the widest single-contract protection — including media liability, production indemnity, loss of licence, legal expenses, management liability and standalone cyber — the TMHCC Media & Music Combined wording offers more covers in one place than any of the five competitor wordings reviewed.", bg=b.BANDBG, border=b.TEAL, lead_color=b.TEAL)
    b.para(doc, "This document is a selection guide only. Cover is subject in all cases to the policy wording, its terms, conditions, exclusions and limits, and to the Schedule.", italic=True, color=b.GREY, size=8.3)

    out = "/app/work/r8/TMHCC_Media_Coverage_Comparison_FULL.docx"
    b.save_doc(doc, out)
    print("D6 Coverage Comparison built:", out)
    print("paras:", len(doc.paragraphs), "tables:", len(doc.tables))


if __name__ == "__main__":
    build()
