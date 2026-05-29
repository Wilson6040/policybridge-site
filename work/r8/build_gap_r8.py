# -*- coding: utf-8 -*-
"""D5 — SEPARATE, INTERNAL Full Gap Analysis (granular).
PART A results + PART B granular gap-by-gap (with copy-paste wording) + PART C sub-limit
harmonisation + Section 12 sign-off list + Reviewer notes / UNCONFIRMED items.
This is the internal/working deficiency & enhancement record — NOT client-facing.
"""
import sys
sys.path.insert(0, '/app/work/compare')
sys.path.insert(0, '/app/work/r8')
from docx import Document
from docx.shared import Pt
import brand2 as b
import cmpdata9 as d
import r8data as r8

PRIO = {"High": ("FBE6EA", "A4233B"), "Medium-High": ("FBF1DA", "8A6A1E"),
        "Medium": ("FBF1DA", "8A6A1E"), "Low-Medium": ("EEF1F3", "5B6166"),
        "Low": ("EEF1F3", "5B6166"), "n/a": ("F2F4F5", "8A9095")}


def field_table(doc, pairs):
    t = b.make_table(doc, ["", ""], [4.4, 13.0])
    # drop header visual by overwriting first row as first datum
    t._tbl.remove(t.rows[0]._tr)
    for k, v in pairs:
        cells = t.add_row().cells
        b.text_cell(cells[0], k, size=8.3, bold=True, color=b.TEAL, bg="EAF1F5")
        b.text_cell(cells[1], v, size=8.4)
    b.spacer(doc, 4)
    return t


def build():
    doc = Document()
    b.set_normal_style(doc)
    b.setup_page(doc, landscape=False, margin_cm=1.7)
    b.add_logo_header(doc, left_text="TMHCC Media & Entertainment — INTERNAL")
    b.add_footer_pagenum(doc, note="INTERNAL / WORKING DOCUMENT — gap analysis & sub-limit record. Not for client distribution. All broadenings subject to TMHCC legal/underwriting sign-off.")

    b.cover_page(
        doc,
        doc_kicker="Gap Analysis & Wording Enhancement — INTERNAL",
        title="Full Gap Analysis",
        subtitle="Granular, clause-for-clause and definition-for-definition deficiency / enhancement record for the TMHCC Media & Music Combined wording, with copy-paste-ready wording for every recommendation and a full sub-limit harmonisation table.",
        meta_lines=[
            "INTERNAL / WORKING DOCUMENT. Separate from the client-facing Coverage Comparison. Every broadening of cover and every sub-limit change is flagged for TMHCC legal / underwriting sign-off.",
            "Grounded in the FINAL TMHCC clean wording and the six attached competitor wordings. Figures that cannot be traced to an attached wording are marked CANNOT-DETERMINE / UNCONFIRMED.",
        ],
        wordings=[f"{w['name']} — {w['full']}" for w in d.WORDINGS],
    )

    doc.add_page_break()
    b.callout(doc, "How to read this document:", "This is the internal gap/enhancement record. PART A reports the specific instructed inclusions (verified then actioned). PART B is the granular gap-by-gap analysis with copy-paste-ready wording. PART C is the sub-limit harmonisation table (old → new). Every item shows priority, underwriting impact, sign-off requirement and whether it is implemented (as a discrete tracked change) or parked. Personal Accident and Travel are deliberately excluded.", bg="FBF1DA", border=b.GOLD, lead_color="8A6A1E")

    # PART A
    b.h1(doc, "Part A — specific master-wording inclusions (verified, then actioned)", num="A")
    for a in r8.PART_A:
        b.h2(doc, f"{a['id']}  ·  {a['title']}")
        field_table(doc, [
            ("Finding", a['finding']),
            ("Location", a['location']),
            ("Action", a['action']),
            ("Status", a['status']),
        ])

    # PART B
    doc.add_page_break()
    b.h1(doc, "Part B — granular gap analysis & bridging", num="B")
    b.para(doc, "Each gap below is classified and, where it is a genuine deficiency supported by an attached competitor wording, bridged in TMHCC house style as a discrete, individually accept/reject-able tracked change. Personal Accident and Travel are excluded by instruction.", align='just')
    for g in r8.GAPS:
        pbg, pfg = PRIO.get(g['priority'], ("F2F4F5", "8A9095"))
        b.h2(doc, f"Gap {g['n']}  ·  {g['title']}")
        field_table(doc, [
            ("Granular item", g['item']),
            ("Current TMHCC position", g['tmhcc']),
            ("Competitor position (cited)", g['comp']),
            ("Recommended action", g['action']),
            ("Exact policy location", g['location']),
            ("Priority", g['priority']),
            ("Underwriting impact", g['uw']),
            ("Legal sign-off required", g['signoff']),
            ("Discrete tracked change applied", g['tracked']),
            ("Status", g['status']),
        ])
        b.callout(doc, "Copy-paste wording:", g['wording'], bg="EEF1F3", border=b.GREY, lead_color=b.TEAL)

    # PART C
    doc.add_page_break()
    b.h1(doc, "Part C — sub-limit harmonisation to market-highest", num="C")
    b.para(doc, "For every extension carrying a sub-limit: TMHCC vs the attached competitor wordings. TMHCC already meets or exceeds every evidenced competitor numeric sub-limit, or is schedule-driven (broker-set). No existing TMHCC sub-limit has been raised because no attached competitor figure is demonstrably higher for the same benefit; none has been invented and none lowered. New gap-fill covers carry sub-limits set to the evidenced market figure.", align='just')
    t = b.make_table(doc, ["Extension / benefit", "Old TMHCC limit", "New TMHCC limit", "Competitor source & limit", "Sign-off"], [4.6, 2.7, 3.0, 5.2, 2.0])
    for ext, old, new, src, so in r8.SUBLIMITS:
        cells = t.add_row().cells
        b.text_cell(cells[0], ext, size=8.1, bold=True, color=b.TEAL)
        b.text_cell(cells[1], old, size=8.0)
        b.text_cell(cells[2], new, size=8.0)
        b.text_cell(cells[3], src, size=8.0)
        b.text_cell(cells[4], so, size=8.0)
    b.zebra(t); b.spacer(doc, 6)

    # Section 12 sign-off list
    b.h1(doc, "Section 12 (Media Liability) — items requiring legal / underwriting sign-off", num="D")
    b.para(doc, "All Section 12 amendments are applied as discrete tracked changes so they can be reviewed and accepted or rejected individually — but each REQUIRES sign-off before acceptance.")
    t = b.make_table(doc, ["Item", "Amendment type", "Sign-off note"], [7.0, 4.5, 6.0])
    for item, typ, note in r8.S12_SIGNOFF:
        cells = t.add_row().cells
        b.text_cell(cells[0], item, size=8.3, bold=True, color=b.TEAL)
        b.text_cell(cells[1], typ, size=8.2)
        b.text_cell(cells[2], note, size=8.2)
    b.zebra(t); b.spacer(doc, 6)

    # Reviewer notes
    b.h1(doc, "Reviewer notes / UNCONFIRMED / cannot-determine", num="E")
    for n in r8.REVIEWER_NOTES:
        b.bullet(doc, n)

    out = "/app/work/r8/TMHCC_Media_GapFill_Enhancement_Strategy.docx"
    b.save_doc(doc, out)
    print("D5 Gap Analysis built:", out)
    print("paras:", len(doc.paragraphs), "tables:", len(doc.tables))


if __name__ == "__main__":
    build()
