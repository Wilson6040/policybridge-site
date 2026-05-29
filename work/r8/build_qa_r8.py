# -*- coding: utf-8 -*-
"""D7 — QA / Change Report."""
import sys
sys.path.insert(0, '/app/work/compare')
sys.path.insert(0, '/app/work/r8')
from docx import Document
from docx.shared import Pt
import brand2 as b
import r8data as r8


def build():
    doc = Document()
    b.set_normal_style(doc)
    b.setup_page(doc, landscape=False, margin_cm=1.7)
    b.add_logo_header(doc, left_text="TMHCC Media & Entertainment — QA")
    b.add_footer_pagenum(doc, note="QA / change report. Internal. Tracked-change wording subject to TMHCC legal / underwriting sign-off.")

    b.cover_page(
        doc,
        doc_kicker="QA / Change Report — Media & Entertainment",
        title="QA & Change Report",
        subtitle="What changed in the TMHCC Media & Music Combined wording, where, what is applied vs parked, and what awaits underwriting/legal sign-off.",
        meta_lines=[
            "Round 8 (EMERGENT DISPATCH v3). Companion to the tracked-changes wording, clean wording, Summary of Changes, Summary of Cover, Gap Analysis and Coverage Comparison.",
        ],
        wordings=["FINAL TMHCC clean wording (baseline)", "Six attached competitor wordings (Tysers, Yutree, Liberty, Allianz, AXA XL, Entertainment Elite)", "Original 0223C wording (Summary of Changes baseline)"],
    )

    doc.add_page_break()
    b.h1(doc, "1  Files produced", num=None)
    for f in [
        "TMHCC_Media_Combined_0526_FINAL_TrackedChanges.docx — final wording with every amendment as a discrete, individually accept/reject-able tracked change (old values visibly replaced).",
        "TMHCC_Media_Combined_0526_FINAL_Clean.docx — clean final wording (all new tracked changes accepted).",
        "TMHCC_Media_Combined_Summary_of_Changes_FINAL.docx (+PDF) — full summary of changes vs the 0223C wording.",
        "TMHCC_Media_Combined_Summary_of_Cover_FINAL.docx (+PDF) — summary of cover, updated to the final wording (no Personal Accident, no Travel).",
        "TMHCC_Media_GapFill_Enhancement_Strategy.docx (+PDF) — SEPARATE internal granular gap analysis & sub-limit record.",
        "TMHCC_Media_Coverage_Comparison_FULL.docx (+PDF) — client/broker-facing coverage comparison (no gap content).",
        "TMHCC_Media_Comparison_QA_Methodology.docx (+PDF) — this QA / change report.",
    ]:
        b.bullet(doc, f)

    b.h1(doc, "2  Document handling")
    for k, v in [
        ("Existing tracked changes accepted before new edits", "Yes — the uploaded clean wording contained no residual revisions; accepted to form the clean baseline."),
        ("Fresh tracked changes used for new edits", "Yes — a new tracked-changes layer was started from the baseline; no old tracked changes carried forward."),
        ("Each amendment discrete / individually accept-reject-able", "Yes — each amendment is a self-contained insertion or deletion at a distinct location."),
        ("Sub-limit / definition changes show old → new", "Yes — definition replacements (A5, A6) delete the old text and insert the new, so the original remains visible."),
        ("Clean final version also produced", "Yes."),
        ("Branding / layout / hyperlinks / numbering", "Preserved — inserted headings reuse existing styles without auto-numbering; no section/column or contents-page changes made."),
    ]:
        b.bullet(doc, v, bold_lead=k + ":  ")

    b.h1(doc, "3  Change log — discrete tracked changes")
    t = b.make_table(doc, ["Ref", "Location", "Type / change", "Old → New", "Sign-off"], [1.3, 4.2, 4.8, 5.7, 1.5])
    for ref, loc, typ, old, new, so in r8.CHANGE_LOG:
        cells = t.add_row().cells
        b.text_cell(cells[0], ref, size=8.2, bold=True, color=b.TEAL)
        b.text_cell(cells[1], loc, size=8.0)
        b.text_cell(cells[2], typ, size=8.0)
        b.text_cell(cells[3], ("OLD: " + old + "\nNEW: " + new), size=7.9)
        b.text_cell(cells[4], so, size=8.0, align='center')
    b.zebra(t); b.spacer(doc, 6)

    b.h1(doc, "4  Part A results")
    t = b.make_table(doc, ["Item", "Result"], [4.0, 13.5])
    for a in r8.PART_A:
        cells = t.add_row().cells
        b.text_cell(cells[0], f"{a['id']} · {a['title']}", size=8.2, bold=True, color=b.TEAL)
        b.text_cell(cells[1], a['status'] + " — " + a['action'], size=8.1)
    b.zebra(t); b.spacer(doc, 6)

    b.h1(doc, "5  Gap & sub-limit summary")
    n_bridged = sum(1 for g in r8.GAPS if g['tracked'] == 'Yes')
    n_parked = sum(1 for g in r8.GAPS if 'PARK' in g['status'].upper())
    n_excl = sum(1 for g in r8.GAPS if 'EXCLUDED' in g['status'].upper())
    n_met = sum(1 for g in r8.GAPS if 'Already met' in g['status'])
    raised = sum(1 for s in r8.SUBLIMITS if 'NEW' in s[0])
    cannot = sum(1 for s in r8.SUBLIMITS if 'cannot determine' in s[4].lower() or 'Cannot determine' in s[3])
    for k, v in [
        ("Gaps reviewed", str(len(r8.GAPS))),
        ("Gaps bridged (discrete tracked change)", str(n_bridged)),
        ("Gaps parked (wording provided, not applied)", str(n_parked)),
        ("Gaps excluded by instruction (Personal Accident & Travel)", str(n_excl)),
        ("Gaps already met (no change)", str(n_met)),
        ("New sub-limited covers added (set to evidenced market figure)", str(raised)),
        ("Existing TMHCC sub-limits raised", "0 — TMHCC already meets/exceeds evidenced figures or is schedule-driven"),
        ("Sub-limits that cannot be determined (no attached figure)", "flagged in Part C; none invented, none lowered"),
    ]:
        b.bullet(doc, v, bold_lead=k + ":  ")

    b.h1(doc, "6  Items awaiting legal / underwriting sign-off (Section 12)")
    for item, typ, note in r8.S12_SIGNOFF:
        b.bullet(doc, f"{typ} — {note}", bold_lead=item + ":  ")

    b.h1(doc, "7  Final QA checklist")
    for c in [
        "Existing tracked changes accepted before new edits — Yes.",
        "New edits shown as fresh, discrete tracked changes — Yes (10 amendments).",
        "Clean final version produced — Yes.",
        "Part A verified-then-actioned: A1 not duplicated; A2 confirmed; A3 not duplicated (nothing lost); A4 touring money added; A5 Employee broadened; A6 cyber definitions aligned.",
        "Gap analysis granular (clause/definition level) and bridges every genuine gap except Personal Accident and Travel.",
        "Sub-limits raised to market-highest where evidenced; un-evidenced flagged, not invented or lowered.",
        "Section 12 amendments flagged for sign-off; cyber definitions aligned §12/§15.",
        "Gap analysis is a SEPARATE document; the Coverage Comparison is client/broker-facing and contains no gap content.",
        "Summary of Changes vs 0223C produced; Summary of Cover updated to final wording only.",
        "No Personal Accident, no Travel added; Personal Assault under Money preserved and distinct.",
        "Branding / layout / hyperlinks / numbering intact.",
        "All UNCONFIRMED / cannot-determine items recorded in Reviewer notes.",
    ]:
        b.bullet(doc, c)

    b.h1(doc, "8  Reviewer notes / UNCONFIRMED")
    for n in r8.REVIEWER_NOTES:
        b.bullet(doc, n)

    out = "/app/work/r8/TMHCC_Media_Comparison_QA_Methodology.docx"
    b.save_doc(doc, out)
    print("D7 QA report built:", out)
    print("paras:", len(doc.paragraphs), "tables:", len(doc.tables))


if __name__ == "__main__":
    build()
