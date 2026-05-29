"""Minimal test: two TOC fields (\\f C and \\f D) in a borderless 2-cell table."""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def fld(fldchar=None, instr=None, text=None):
    r=OxmlElement('w:r')
    if fldchar:
        fc=OxmlElement('w:fldChar');fc.set(qn('w:fldCharType'),fldchar);r.append(fc)
    if instr is not None:
        it=OxmlElement('w:instrText');it.set(qn('xml:space'),'preserve');it.text=instr;r.append(it)
    if text is not None:
        t=OxmlElement('w:t');t.set(qn('xml:space'),'preserve');t.text=text;r.append(t)
    return r

def add_tc(p,text,ident):
    pe=p._p
    pe.append(fld(fldchar='begin')); pe.append(fld(instr=f' TC "{text}" \\f {ident} \\l "1" ')); pe.append(fld(fldchar='end'))

def toc_into_cell(cell, ident):
    p=cell.paragraphs[0]
    pe=p._p
    pe.append(fld(fldchar='begin')); pe.append(fld(instr=f' TOC \\f {ident} \\h \\z ')); pe.append(fld(fldchar='separate')); pe.append(fld(text='update')); pe.append(fld(fldchar='end'))

doc=Document()
doc.add_paragraph("Policy Contents")
table=doc.add_table(rows=1,cols=2)
# borderless
tblPr=table._tbl.tblPr
borders=OxmlElement('w:tblBorders')
for e in ('top','left','bottom','right','insideH','insideV'):
    el=OxmlElement('w:'+e);el.set(qn('w:val'),'none');borders.append(el)
tblPr.append(borders)
w=OxmlElement('w:tblW');w.set(qn('w:type'),'pct');w.set(qn('w:w'),'5000');tblPr.append(w)
toc_into_cell(table.rows[0].cells[0],'C')
toc_into_cell(table.rows[0].cells[1],'D')
doc.add_page_break()
for i in range(1,4):
    p=doc.add_paragraph(f"First-half Heading {i}"); add_tc(p,f"First-half Heading {i}",'C')
    doc.add_paragraph("body "*300); doc.add_page_break()
for i in range(1,4):
    p=doc.add_paragraph(f"Second-half Heading {i}"); add_tc(p,f"Second-half Heading {i}",'D')
    doc.add_paragraph("body "*300); doc.add_page_break()
uf=OxmlElement('w:updateFields');uf.set(qn('w:val'),'true');doc.settings.element.append(uf)
doc.save('/app/work/toc_table_test.docx')
print("saved")
