# -*- coding: utf-8 -*-
import sys
sys.path.insert(0, '/app/work/compare')
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import brand2 as b
import cmpdata as d

PRIO = {
    "High":        ("FBE6EA", "A4233B"),
    "Medium-High": ("FBE6EA", "A4233B"),
    "Medium":      ("FBF1DA", "8A6A1E"),
    "Low-Medium":  ("EEF1F3", "5B6166"),
    "Low":         ("EEF1F3", "5B6166"),
}


def prio_chip(doc, priority):
    bg, fg = PRIO.get(priority, ("EEF1F3", "5B6166"))
    p = doc.paragraphs[-1]
    r = p.add_run("   [ Priority: " + priority + " ]")
    b._set_run(r, size=9, bold=True, color=fg)


def attr_table(doc, pairs, label_w=4.3, val_w=13.4):
    t = doc.add_table(rows=0, cols=2)
    b._table_borders(t, b.LINE, 4)
    b._table_full_width(t)
    b._set_col_widths(t, [label_w, val_w])
    for k, v in pairs:
        cells = t.add_row().cells
        b.text_cell(cells[0], k, size=8.6, bold=True, color=b.TEAL, bg="EAF1F5")
        b.text_cell(cells[1], v, size=8.8)
    b.spacer(doc, 8)
    return t


def rec_card(doc, rec, idx):
    b.h2(doc, f"{idx}.  {rec['title']}")
    prio_chip(doc, rec['priority'])
    attr_table(doc, [
        ("Competitor(s)", rec['comp']),
        ("What the competitor does", rec['does']),
        ("Current TMHCC position", rec['tmhcc']),
        ("Gap / opportunity", rec['gap']),
        ("Proposed TMHCC action", rec['action']),
        ("Commercial benefit", rec['commercial']),
        ("Broker / client benefit", rec['broker']),
        ("Underwriting risk", rec['uw']),
        ("Claims risk", rec['claims']),
        ("Legal / underwriting review", rec['legal']),
    ])


def build():
    doc = Document()
    b.set_normal_style(doc)
    b.setup_page(doc, landscape=False, margin_cm=1.7)
    b.add_logo_header(doc)
    b.add_footer_pagenum(doc)

    # ---- COVER ----
    b.cover_page(
        doc,
        doc_kicker="Wording Enhancement \u2014 Media & Entertainment",
        title="TMHCC Gap-Fill &\nWording Enhancement Strategy",
        subtitle="Strategic, action-focused recommendations to make the TMHCC Media & Music Combined base wording more holistic, clearer and demonstrably market-leading.",
        meta_lines=[
            "Companion to the Full Coverage Comparison. Every recommendation is evidence-based \u2014 supported by at least one competitor wording.",
            "Recommendations are commercial proposals only and are subject to final TMHCC legal / underwriting sign-off.",
        ],
        wordings=[
            "Benchmarked against: Tysers (Zurich), Yutree (AXA), Liberty, Allianz",
            "Fifth competitor (Wording E): pending \u2014 to be incorporated later",
        ],
    )

    # ---- EXEC SUMMARY ----
    doc.add_page_break()
    b.h1(doc, "Executive summary", num="1")
    b.para(doc, "The Full Coverage Comparison confirms that TMHCC already offers the broadest media & entertainment package in the peer group \u2014 the only wording with all fifteen sections, and the only one with Loss of Licence, Legal Expenses, Management Liability and a genuine standalone Cyber section. The purpose of this document is not to broaden recklessly, but to close the small number of places where a competitor genuinely shows more, so that TMHCC can present a tick on virtually every comparison row.", align='just')
    b.h3(doc, "What moves the needle most")
    b.bullet(doc, "Add a Personal Accident & Business Travel option \u2014 the only whole section two competitors offer that TMHCC does not.", bold_lead="1.  ")
    b.bullet(doc, "Offer a rated worldwide / US-Canada territory option for Media Liability \u2014 Tysers PI is worldwide.", bold_lead="2.  ")
    b.bullet(doc, "Add a Distributors & Purchasers extension and journalistic source-protection costs to Media Liability (both Tysers).", bold_lead="3.  ")
    b.bullet(doc, "Clarify where criminal / regulatory defence costs live (S12 vs S13/S14) \u2014 a near nil-cost fix that removes an apparent gap.", bold_lead="4.  ")
    b.callout(doc, "Discipline retained:", "TMHCC should KEEP its more restrictive positions on patents, communicable disease and default US/Canada jurisdiction. These are deliberate, defensible underwriting choices, not gaps \u2014 the worldwide/US-Canada exposure is best offered as a rated option rather than given away.")

    # ---- STRATEGIC OBJECTIVE ----
    doc.add_page_break()
    b.h1(doc, "Strategic objective", num="2")
    b.para(doc, "To evolve the TMHCC base wording from \u2018already the broadest\u2019 to \u2018demonstrably market-leading on every row\u2019 \u2014 by:")
    b.bullet(doc, "filling the few genuine coverage gaps a competitor exposes;")
    b.bullet(doc, "adding targeted exclusion write-backs where they are commercially useful and underwriting-acceptable;")
    b.bullet(doc, "sharpening definitions and conditions for clarity; and")
    b.bullet(doc, "doing all of the above without diluting TMHCC\u2019s disciplined risk selection or its protective claims architecture.")
    b.para(doc, "Each recommendation below states the competitor evidence, the current TMHCC position, the proposed action and the underwriting/claims impact, with a priority rating.", space_before=4)

    # ---- PRIORITY SUMMARY ----
    doc.add_page_break()
    b.h1(doc, "Priority gap-fill recommendations", num="3")
    b.para(doc, "All recommendations at a glance, ordered by priority.")
    t = b.make_table(doc, ["#", "Recommendation", "Evidence (competitor)", "Action type", "Priority"],
                     [0.8, 6.9, 4.2, 3.3, 2.6])
    action_type = {
        0: "Add section/extension", 1: "Add rated option", 2: "Add extension",
        3: "Clarify / cross-ref", 4: "Add extension", 5: "Add extension",
        6: "Add write-back", 7: "Amend condition", 8: "Clarify wording",
        9: "Clarify wording", 10: "Amend definition", 11: "Leave unchanged",
    }
    order = sorted(range(len(d.GAPFILL)), key=lambda i: ["High","Medium-High","Medium","Low-Medium","Low"].index(d.GAPFILL[i]['priority']))
    for n, i in enumerate(order, start=1):
        rec = d.GAPFILL[i]
        cells = t.add_row().cells
        b.text_cell(cells[0], str(n), size=8.5, bold=True, color=b.TEAL, align='center')
        b.text_cell(cells[1], rec['title'], size=8.3, bold=True, color=b.INK)
        b.text_cell(cells[2], rec['comp'], size=8.0)
        b.text_cell(cells[3], action_type.get(i, ""), size=8.0)
        bg, fg = PRIO.get(rec['priority'], ("EEF1F3", "5B6166"))
        b.text_cell(cells[4], rec['priority'], size=8.2, bold=True, color=fg, bg=bg, align='center')
    b.zebra(t)

    # ---- 4. COVERAGE GAPS TO CONSIDER FILLING (cards) ----
    doc.add_page_break()
    b.h1(doc, "Coverage gaps to consider filling", num="4")
    b.para(doc, "Detailed recommendation cards, in priority order.")
    for n, i in enumerate(order, start=1):
        rec_card(doc, d.GAPFILL[i], n)

    # ---- 5. EXCLUSION WRITE-BACKS ----
    doc.add_page_break()
    b.h1(doc, "Exclusion write-backs to consider", num="5")
    b.para(doc, "Where a competitor narrows an exclusion or writes cover back, and whether TMHCC should follow. Exclusions are not softened lightly \u2014 each note states the underwriting view.")
    t = b.make_table(doc, ["Exclusion area", "Competitor", "What they do", "TMHCC position", "Recommended TMHCC action", "Priority"],
                     [3.1, 2.2, 3.3, 3.2, 4.0, 2.0])
    for area, comp, does, pos, action, prio in d.EXCL_WRITEBACKS:
        cells = t.add_row().cells
        b.text_cell(cells[0], area, size=8.0, bold=True, color=b.TEAL)
        b.text_cell(cells[1], comp, size=7.9)
        b.text_cell(cells[2], does, size=7.9)
        b.text_cell(cells[3], pos, size=7.9)
        b.text_cell(cells[4], action, size=7.9)
        bg, fg = PRIO.get(prio, ("EEF1F3", "5B6166"))
        b.text_cell(cells[5], prio, size=7.9, bold=True, color=fg, bg=bg, align='center')
    b.zebra(t)
    b.callout(doc, "Underwriting stance:", "Adopt the asbestos and pollution write-backs ONLY in a narrow, sub-limited, negligent-advice form that keeps bodily injury and the contamination/clean-up itself excluded. Offer worldwide/US-Canada as a rated option. Do NOT write patents back beyond, at most, limited defence costs.")

    # ---- 6. DEFINITIONS ----
    doc.add_page_break()
    b.h1(doc, "Definitions to consider expanding or clarifying", num="6")
    defs = [
        ("\u2018Media Material\u2019 / \u2018Content\u2019", "Tysers defines \u2018Media Material\u2019",
         "TMHCC uses \u2018data, text, sounds, images or similar content\u2019 descriptively.",
         "Add a defined term to sharpen the subject-matter of S12 cover.", "Low"),
        ("Territory \u2014 \u2018Geographical Limits\u2019 / \u2018Jurisdiction\u2019", "Tysers (worldwide)",
         "Default worldwide ex-US/Canada; Legal Action exclusion bars actions outside Jurisdiction.",
         "Define a clear optional worldwide / incl.-US-Canada basis selectable in the Schedule.", "High"),
        ("\u2018Cyber act\u2019 / \u2018Computer System\u2019 / \u2018Virus\u2019", "Yutree defines \u2018Cyber act\u2019",
         "TMHCC defines \u2018Computer System\u2019 and \u2018Virus\u2019 within S12 and operates a separate CyberGuard section.",
         "Align S12 cyber definitions with the CyberGuard section to avoid overlap/ambiguity.", "Medium"),
        ("\u2018Employee\u2019 / freelancers & contractors", "Tysers; Yutree (broad employee defs)",
         "TMHCC S12 includes self-employed/freelance under direction & control \u2014 already broad.",
         "Confirm freelancers/loaned staff captured consistently across sections.", "Low"),
    ]
    t = b.make_table(doc, ["Definition", "Competitor benchmark", "Current TMHCC", "Proposed clarification", "Priority"],
                     [3.4, 3.2, 3.8, 5.0, 2.4])
    for name, bench, cur, prop, prio in defs:
        cells = t.add_row().cells
        b.text_cell(cells[0], name, size=8.0, bold=True, color=b.TEAL)
        b.text_cell(cells[1], bench, size=7.9)
        b.text_cell(cells[2], cur, size=7.9)
        b.text_cell(cells[3], prop, size=7.9)
        bg, fg = PRIO.get(prio, ("EEF1F3", "5B6166"))
        b.text_cell(cells[4], prio, size=7.9, bold=True, color=fg, bg=bg, align='center')
    b.zebra(t)

    # ---- 7. CONDITIONS / CLAIMS OBLIGATIONS ----
    doc.add_page_break()
    b.h1(doc, "Conditions / claims obligations to consider improving", num="7")
    b.bullet(doc, "Adopt a King\u2019s Counsel dispute-resolution clause in S12 (as Yutree) to give clients certainty on defence/settlement disputes \u2014 procedural, low risk.", bold_lead="Adopt:  ")
    b.bullet(doc, "Add an express cross-reference in S12 confirming that statutory / criminal-defence cover is provided under S13 (Legal Expenses) and S14 (Management Liability).", bold_lead="Clarify:  ")
    b.bullet(doc, "Add a signpost in S1/S2 confirming computer / IT breakdown is included (Yutree shows a named Computer Breakdown section).", bold_lead="Clarify:  ")
    b.bullet(doc, "RETAIN claim-notification as an express condition precedent \u2014 this is more protective than the competitors and should not be softened.", bold_lead="Retain:  ")
    b.bullet(doc, "RETAIN the detailed S12 computer-security conditions \u2014 strong risk control; ensure they are workable for SME media clients.", bold_lead="Retain:  ")

    # ---- 8. LIMITS / SUB-LIMITS ----
    doc.add_page_break()
    b.h1(doc, "Limits and sub-limits to consider reviewing", num="8")
    t = b.make_table(doc, ["Sub-limit", "TMHCC today", "Competitor benchmark", "Recommendation"],
                     [4.0, 3.6, 4.4, 5.8])
    rows = [
        ("Criminal / statutory defence (S12)", "DP defence only (GBP 250k)", "Tysers GBP 1m; Yutree GBP 250k",
         "Add a modest S12 criminal/regulatory defence sub-limit OR cross-refer to S13/S14."),
        ("Representation costs (S12)", "Via DP defence only", "Tysers GBP 25k",
         "Introduce a GBP 25k representation-costs sub-limit."),
        ("Distributors & purchasers (S12)", "Not offered", "Tysers up to 5x limit",
         "If adopted, cap the aggregate (e.g. a defined multiple of the limit)."),
        ("Loss of documents (S12)", "Per Schedule", "Tysers GBP 1m",
         "Confirm an explicit sub-limit for parity of presentation."),
        ("Virus transmission (S12)", "GBP 500k", "Tysers GBP 250k",
         "TMHCC already higher \u2014 retain as a strength."),
    ]
    for r in rows:
        cells = t.add_row().cells
        b.text_cell(cells[0], r[0], size=8.1, bold=True, color=b.TEAL)
        b.text_cell(cells[1], r[1], size=8.0)
        b.text_cell(cells[2], r[2], size=8.0)
        b.text_cell(cells[3], r[3], size=8.0)
    b.zebra(t)

    # ---- 9. MARKET-LEADING OPPORTUNITIES ----
    doc.add_page_break()
    b.h1(doc, "Market-leading wording opportunities", num="9")
    b.para(doc, "Implemented together, the high-priority items let TMHCC claim a genuine, evidence-based market-leading position:")
    for s in [
        "\u2018The only media & entertainment wording with all fifteen sections \u2014 now including a Personal Accident & Business Travel option.\u2019",
        "\u2018Worldwide media-liability cover available, including a rated US/Canada extension.\u2019",
        "\u2018Distribution chain protected \u2014 distributors & purchasers extension.\u2019",
        "\u2018Standalone cyber, legal expenses and management liability built in \u2014 not bolt-ons.\u2019",
        "\u2018Reputation management, withdrawal-of-content and data-protection defence as standard in media liability.\u2019",
    ]:
        b.bullet(doc, s)
    b.callout(doc, "Positioning:", "TMHCC can credibly market the wording as broader, clearer and more complete than the four competitors reviewed \u2014 subject to legal/underwriting sign-off and to incorporating the fifth competitor when supplied.")

    # ---- 10. UNDERWRITING RISK ASSESSMENT ----
    doc.add_page_break()
    b.h1(doc, "Underwriting risk assessment", num="10")
    t = b.make_table(doc, ["Recommendation", "UW risk", "Control / stance", "Verdict"],
                     [6.0, 2.4, 6.4, 3.0])
    risk_rows = [
        ("Personal Accident & Business Travel", "Medium", "Benefit schedule, age limits, hazardous-activity carve-outs", "Adopt (capacity-permitting)"),
        ("Worldwide / US-Canada PI option", "High (if US/Canada)", "Rated option, sub-limit, higher excess, defence caps; keep ex-US/Canada default", "Adopt as option only"),
        ("Distributors & purchasers", "Medium", "36-month window, certificate, control of defence, capped aggregate", "Adopt with controls"),
        ("Criminal-defence clarity / cross-ref", "Low", "Cross-reference to S13/S14; small S12 sub-limit", "Adopt"),
        ("Source-protection costs", "Low", "Small sub-limit; prospects test; consent", "Adopt"),
        ("Representation costs", "Low", "GBP 25k sub-limit", "Adopt"),
        ("Asbestos / pollution PI write-back", "Medium", "Negligent-advice only; exclude BI & contamination; tight cap", "Adopt narrowly"),
        ("Patents", "High", "Systemic severity", "Leave excluded (defence-costs only at most)"),
        ("Communicable disease / US-Canada default", "High", "Retain broad exclusion / default territory", "Leave unchanged"),
    ]
    for r in risk_rows:
        cells = t.add_row().cells
        b.text_cell(cells[0], r[0], size=8.1, bold=True, color=b.TEAL)
        b.text_cell(cells[1], r[1], size=8.0, align='center')
        b.text_cell(cells[2], r[2], size=8.0)
        b.text_cell(cells[3], r[3], size=8.0, bold=True)
    b.zebra(t)

    # ---- 11. IMPLEMENTATION ROADMAP ----
    doc.add_page_break()
    b.h1(doc, "Implementation roadmap", num="11")
    phases = [
        ("Phase 1 \u2014 Clarifications (nil / low cost)",
         ["Cross-reference criminal/regulatory defence to S13/S14.",
          "Signpost computer/IT breakdown within S1/S2.",
          "Name ICOW/AICOW and book-debts heads expressly in S3.",
          "Add \u2018Media Material\u2019 definition; align cyber definitions."]),
        ("Phase 2 \u2014 Sub-limited extensions (low risk)",
         ["Add representation-costs (GBP 25k) and source-protection costs to S12.",
          "Adopt the King\u2019s Counsel dispute-resolution clause.",
          "Confirm explicit loss-of-documents sub-limit."]),
        ("Phase 3 \u2014 Rated options & new cover (priced)",
         ["Add Personal Accident & Business Travel option.",
          "Add rated worldwide / US-Canada PI territory option.",
          "Add Distributors & Purchasers extension with capped aggregate.",
          "Consider narrow asbestos / pollution negligent-advice write-backs."]),
        ("Phase 4 \u2014 Document deliberate restrictions",
         ["Record the rationale for retaining the patent, communicable-disease and default US/Canada exclusions.",
          "Re-run the comparison once Wording E (fifth competitor) is supplied."]),
    ]
    for title, items in phases:
        b.h3(doc, title)
        for it in items:
            b.bullet(doc, it)
        b.spacer(doc, 2)

    # ---- 12. LEGAL / UW SIGN-OFF ----
    doc.add_page_break()
    b.h1(doc, "Legal / underwriting sign-off notes", num="12")
    for n in [
        "All recommendations are commercial proposals and require TMHCC legal and underwriting approval before adoption.",
        "Any write-back must be drafted to preserve the existing exclusion\u2019s intent (especially bodily injury, contamination and systemic risk).",
        "Rated options (PA, worldwide/US-Canada, distributors) require pricing, reinsurance-treaty and capacity confirmation.",
        "Confirm freelancer/contractor and \u2018insured\u2019 definitions remain consistent across all amended sections.",
        "No cover has been invented in this strategy; every item is supported by at least one competitor wording cited above.",
        "Incorporate the fifth competitor (Wording E) and revisit priorities once that wording is supplied.",
    ]:
        b.bullet(doc, n)
    b.spacer(doc, 6)
    b.para(doc, "End of Gap-Fill / Wording Enhancement Strategy.", italic=True, color=b.GREY, size=8.5)

    out = "/app/work/compare/out/TMHCC_Media_GapFill_Enhancement_Strategy.docx"
    b.save_doc(doc, out)
    print("Saved:", out)
    return out


if __name__ == "__main__":
    build()
