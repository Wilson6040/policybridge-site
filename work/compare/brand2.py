"""
Branding helpers for the TMHCC Media & Entertainment coverage-comparison
deliverables. Replicates the supplied "Coverage Comparison" template:
 - Tokio Marine HCC logo (image1.png) top-right
 - Lato font (Arial fallback), Normal 9.5pt #2B2F33
 - Palette: teal #00648B (headings), gold #B88A3C (subtitle), rust #C0563F (accent)
 - Legend  : ✓ Covered · ◐ Partial/conditional/sub-limit · ✗ Not covered · ? Requires review
 - Clean status tables with coloured chips, brand header band, footer page numbers.
"""
import os, zipfile, subprocess, shutil
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT

# ---------------- Brand palette ----------------
TEAL   = "00648B"   # primary brand blue/teal (headings, table header band)
TEAL_D = "004A66"   # darker teal
GOLD   = "B88A3C"   # subtitle / accent gold-bronze
RUST   = "C0563F"   # accent rust/red
GREY   = "6C7378"   # legend grey
LGREY  = "AEB4B8"   # meta light grey
INK    = "2B2F33"   # body text (dark slate)
LINE   = "D6DCE0"   # table border light grey
BANDBG = "EAF1F5"   # very light teal band
ZEBRA  = "F4F7F9"   # zebra row tint

FONT   = "Lato"

LOGO   = "/app/work/compare/tmhcc_logo.png"

# Status chips: key -> (symbol, bg, fg, label)
STATUS = {
    "yes":     ("\u2713", "E3F1E8", "1B7A3D", "Covered"),
    "partial": ("\u25D0", "FBF1DA", "8A6A1E", "Partial / conditional"),
    "no":      ("\u2717", "FBE6EA", "A4233B", "Not covered"),
    "review":  ("?",      "EEF1F3", "5B6166", "Requires review"),
    "na":      ("\u2014", "F2F4F5", "8A9095", "Not applicable"),
    "pending": ("\u2026", "F2F4F5", "8A9095", "Pending"),
}
# comparison verdict chips
VERDICT = {
    "tmhcc":   ("D9EAF2", "00648B"),   # TMHCC stronger / broader / clearer
    "comp":    ("FBE6EA", "A4233B"),   # competitor stronger / broader
    "equal":   ("EFF1F2", "5B6166"),   # broadly equivalent
    "review":  ("FBF1DA", "8A6A1E"),   # unclear / requires review
}


# ---------------- low-level run / font ----------------
def _set_run(run, size=9.5, bold=False, italic=False, color=INK, font=FONT, caps=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), font)
    if caps:
        c = OxmlElement('w:caps'); c.set(qn('w:val'), 'true'); rpr.append(c)


def set_normal_style(doc):
    st = doc.styles['Normal']
    st.font.name = FONT
    st.font.size = Pt(9.5)
    st.font.color.rgb = RGBColor.from_string(INK)
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), FONT)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.08


# ---------------- page setup ----------------
def setup_page(doc, landscape=False, margin_cm=1.6):
    sec = doc.sections[0]
    if landscape:
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width  = Cm(29.7); sec.page_height = Cm(21.0)
    else:
        sec.orientation = WD_ORIENT.PORTRAIT
        sec.page_width  = Cm(21.0); sec.page_height = Cm(29.7)
    for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
        setattr(sec, m, Cm(margin_cm))
    sec.header_distance = Cm(0.8); sec.footer_distance = Cm(0.8)
    return sec


def add_logo_header(doc, left_text="TMHCC Media & Entertainment"):
    """Header: small brand label left, Tokio Marine HCC logo right.
    Width-aware so it fits both portrait and landscape pages."""
    sec = doc.sections[0]
    sec.different_first_page_header_footer = False
    # usable width = page width - margins (EMU)
    usable = sec.page_width - sec.left_margin - sec.right_margin
    hdr = sec.header
    hdr.is_linked_to_previous = False
    # use a 2-col table so label sits left and logo sits right
    p = hdr.paragraphs[0]
    p.text = ""
    tbl = hdr.add_table(rows=1, cols=2, width=usable)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_full_width(tbl)
    c0, c1 = tbl.rows[0].cells
    c0.width = Emu(int(usable * 0.74)); c1.width = Emu(int(usable * 0.26))
    c0.vertical_alignment = 1
    pp = c0.paragraphs[0]
    r = pp.add_run(left_text)
    _set_run(r, size=8, bold=True, color=GREY)
    pp.paragraph_format.space_after = Pt(0)
    pc = c1.paragraphs[0]
    pc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = pc.add_run()
    try:
        run.add_picture(LOGO, height=Cm(0.85))
    except Exception:
        pass
    _no_table_borders(tbl)
    # bottom rule line under header
    return hdr


def add_footer_pagenum(doc, note="Indicative wording comparison — broking reference only. Refer to the full policy wordings. Subject to TMHCC legal/underwriting sign-off."):
    sec = doc.sections[0]
    ftr = sec.footer
    ftr.is_linked_to_previous = False
    p = ftr.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(note)
    _set_run(r, size=6.5, color=LGREY, italic=True)
    # page number paragraph (centred)
    p2 = ftr.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_field(p2)


def _add_page_field(p):
    r = p.add_run(); _set_run(r, size=7.5, color=GREY)
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = ' PAGE '
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    r._r.append(fld1); r._r.append(instr); r._r.append(fld2)


# ---------------- shading / borders ----------------
def _shade(el, hexcolor):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    el.append(shd)


def _cell_bg(cell, hexcolor):
    _shade(cell._tc.get_or_add_tcPr(), hexcolor)


def _no_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge); e.set(qn('w:val'), 'none'); borders.append(e)
    tblPr.append(borders)


def _table_borders(table, color=LINE, sz=4):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        borders.append(e)
    tblPr.append(borders)


def _table_full_width(table):
    tblPr = table._tbl.tblPr
    w = OxmlElement('w:tblW'); w.set(qn('w:type'), 'pct'); w.set(qn('w:w'), '5000')
    tblPr.append(w)
    table.autofit = False
    # fixed layout for predictable column widths
    lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'), 'fixed'); tblPr.append(lay)


def _set_col_widths(table, widths_cm):
    """Set explicit grid + cell widths in twips and pin the table to the
    summed width with a fixed layout, so it can never exceed the page."""
    EMU_PER_CM = 360000
    TWIPS_PER_CM = 567
    total_tw = int(round(sum(widths_cm) * TWIPS_PER_CM))
    tbl = table._tbl
    tblPr = tbl.tblPr
    # table total width (dxa) + fixed layout
    for tag in ('w:tblW', 'w:tblLayout'):
        ex = tblPr.find(qn(tag))
        if ex is not None:
            tblPr.remove(ex)
    w = OxmlElement('w:tblW'); w.set(qn('w:type'), 'dxa'); w.set(qn('w:w'), str(total_tw))
    tblPr.append(w)
    lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'), 'fixed'); tblPr.append(lay)
    table.autofit = False
    # rebuild tblGrid
    grid = tbl.find(qn('w:tblGrid'))
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement('w:tblGrid')
    for wcm in widths_cm:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(round(wcm * TWIPS_PER_CM))))
        grid.append(gc)
    tblPr.addnext(grid)
    # per-cell widths
    for j, wcm in enumerate(widths_cm):
        for row in table.rows:
            if j < len(row.cells):
                row.cells[j].width = Cm(wcm)


# ---------------- text helpers ----------------
def para(doc, text="", size=9.5, bold=False, italic=False, color=INK, align=None,
         space_after=4, space_before=0, font=FONT):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if align == 'right':  p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if align == 'just':   p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if text:
        r = p.add_run(text); _set_run(r, size=size, bold=bold, italic=italic, color=color, font=font)
    return p


def rich(doc, parts, size=9.5, align=None, space_after=4, space_before=0):
    """parts: list of (text, {bold,italic,color}) tuples."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for txt, opt in parts:
        r = p.add_run(txt)
        _set_run(r, size=size, bold=opt.get('bold', False), italic=opt.get('italic', False),
                 color=opt.get('color', INK))
    return p


def h1(doc, text, num=None):
    """Teal section heading with a gold rule beneath."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    if num:
        r = p.add_run(f"{num}  ")
        _set_run(r, size=15, bold=True, color=GOLD)
    r = p.add_run(text)
    _set_run(r, size=15, bold=True, color=TEAL)
    _bottom_border(p, GOLD, 8)
    return p


def h2(doc, text, color=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text); _set_run(r, size=11.5, bold=True, color=color)
    return p


def h3(doc, text, color=GOLD):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text); _set_run(r, size=10, bold=True, color=color)
    return p


def bullet(doc, text, level=0, bold_lead=None, color=INK, size=9.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.6 + 0.5 * level)
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r = p.add_run(bold_lead); _set_run(r, size=size, bold=True, color=TEAL)
    r = p.add_run(text); _set_run(r, size=size, color=color)
    return p


def _bottom_border(p, color, sz=6):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '2'); bottom.set(qn('w:color'), color)
    pbdr.append(bottom); pPr.append(pbdr)


def spacer(doc, pts=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(pts); return p


# ---------------- legend ----------------
def legend_bar(doc):
    items = [("\u2713", "1B7A3D", "Covered"),
             ("\u25D0", "8A6A1E", "Partial / conditional / sub-limit"),
             ("\u2717", "A4233B", "Not covered / not identified"),
             ("?", "5B6166", "Unclear \u2013 requires review")]
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
    for i, (sym, col, lbl) in enumerate(items):
        if i:
            sep = p.add_run("     "); _set_run(sep, size=8.5, color=LGREY)
        r = p.add_run(sym + " "); _set_run(r, size=9.5, bold=True, color=col)
        r2 = p.add_run(lbl); _set_run(r2, size=8.5, color=GREY)
    return p


# ---------------- status table ----------------
def status_cell(cell, key, small=False):
    info = STATUS.get(key)
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    if info:
        sym, bg, fg, lbl = info
        _cell_bg(cell, bg)
        r = p.add_run(sym); _set_run(r, size=10.5 if not small else 9.5, bold=True, color=fg)
    else:
        r = p.add_run(str(key)); _set_run(r, size=8.5, color=INK)


def text_cell(cell, text, size=8.5, bold=False, color=INK, bg=None, align=None, caps=False):
    cell.text = ""
    if bg: _cell_bg(cell, bg)
    parts = str(text).split("\n")
    for i, part in enumerate(parts):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(0)
        if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(part); _set_run(r, size=size, bold=bold, color=color, caps=caps)


def header_row(table, headers, bg=TEAL, fg="FFFFFF", size=8.5):
    for j, h in enumerate(headers):
        c = table.rows[0].cells[j]
        text_cell(c, h, size=size, bold=True, color=fg, bg=bg, align='center')


def make_table(doc, headers, widths_cm, header_bg=TEAL):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_borders(table, LINE, 4)
    header_row(table, headers, bg=header_bg)
    if widths_cm:
        _set_col_widths(table, widths_cm)
    return table


def zebra(table, start=1):
    for i, row in enumerate(table.rows):
        if i >= start and (i - start) % 2 == 1:
            for c in row.cells:
                if c._tc.get_or_add_tcPr().find(qn('w:shd')) is None:
                    _cell_bg(c, ZEBRA)


# ---------------- callout ----------------
def callout(doc, lead, text, bg=BANDBG, border=TEAL, lead_color=TEAL):
    table = doc.add_table(rows=1, cols=1)
    _table_full_width(table)
    cell = table.rows[0].cells[0]
    _cell_bg(cell, bg)
    tcPr = cell._tc.get_or_add_tcPr()
    bdr = OxmlElement('w:tcBorders')
    for edge, c, s in (('top', border, 4), ('left', border, 22), ('bottom', border, 4), ('right', border, 4)):
        e = OxmlElement('w:' + edge); e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(s)); e.set(qn('w:space'), '0'); e.set(qn('w:color'), c)
        bdr.append(e)
    tcPr.append(bdr)
    cell.text = ""
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    if lead:
        r = p.add_run(lead + "  "); _set_run(r, size=9.5, bold=True, color=lead_color)
    r = p.add_run(text); _set_run(r, size=9.5, color=INK)
    spacer(doc, 4)
    return table


# ---------------- cover page ----------------
def cover_page(doc, doc_kicker, title, subtitle, meta_lines, wordings):
    sec = doc.sections[0]
    # logo top
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    try: run.add_picture(LOGO, height=Cm(1.5))
    except Exception: pass
    spacer(doc, 30)
    # kicker (gold caps)
    para(doc, doc_kicker, size=11, bold=True, color=GOLD, caps_hack=False) if False else None
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(doc_kicker); _set_run(r, size=11, bold=True, color=GOLD, caps=True)
    # title (big teal)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title); _set_run(r, size=30, bold=True, color=TEAL)
    _bottom_border(p, GOLD, 14)
    # subtitle
    para(doc, subtitle, size=13, bold=False, color=INK, space_before=6, space_after=14)
    # wordings reviewed band
    h3(doc, "Wordings reviewed")
    for w in wordings:
        bullet(doc, w)
    spacer(doc, 10)
    for m in meta_lines:
        para(doc, m, size=8.5, italic=True, color=GREY, space_after=3)


# patch para() to ignore stray kw
_orig_para = para
def para(doc, text="", size=9.5, bold=False, italic=False, color=INK, align=None,
         space_after=4, space_before=0, font=FONT, **kw):
    return _orig_para(doc, text=text, size=size, bold=bold, italic=italic, color=color,
                      align=align, space_after=space_after, space_before=space_before, font=font)


# ---------------- save + PDF ----------------
def save_doc(doc, path):
    doc.save(path)
    return path


def to_pdf(docx_path, out_dir):
    os.makedirs(out_dir, exist_ok=True)
    subprocess.run(['pkill', '-f', 'soffice'], capture_output=True)
    import time; time.sleep(1)
    r = subprocess.run(
        ['soffice', '--headless', '--convert-to',
         'pdf:writer_pdf_Export', '--outdir', out_dir, docx_path],
        capture_output=True, text=True, timeout=180)
    pdf = os.path.join(out_dir, os.path.splitext(os.path.basename(docx_path))[0] + '.pdf')
    return pdf if os.path.exists(pdf) else None
