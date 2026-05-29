# -*- coding: utf-8 -*-
import sys
sys.path.insert(0, '/app/work/compare')
from docx import Document
from docx.shared import Cm, Pt
import brand2 as b
import cmpdata as d

VERDICT_LABEL = {
    "tmhcc": ("TMHCC stronger", b.TEAL),
    "comp":  ("Competitor broader", "A4233B"),
    "equal": ("Equivalent", "5B6166"),
    "review":("Review", "8A6A1E"),
}

W_HEADERS = ["TMHCC", "Tysers\n(Zurich)", "Yutree\n(AXA)", "Liberty", "Allianz", "Wording E\n(pending)"]


def comment_cell(cell, verdict, text):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    if verdict and verdict in VERDICT_LABEL:
        lbl, col = VERDICT_LABEL[verdict]
        r = p.add_run(lbl + " \u00b7 ")
        b._set_run(r, size=8, bold=True, color=col)
    r = p.add_run(text)
    b._set_run(r, size=8, color=b.INK)


def render_matrix(doc, label_head, rows, label_w=5.4, comment_w=11.6, label_col0_teal=True):
    headers = [label_head] + W_HEADERS + ["Comment / practical impact"]
    status_w = (26.0 - label_w - comment_w) / 6.0
    widths = [label_w] + [status_w] * 6 + [comment_w]
    t = b.make_table(doc, headers, widths)
    for row in rows:
        label, statuses, verdict, comment = row
        cells = t.add_row().cells
        b.text_cell(cells[0], label, size=8.3, bold=True, color=b.TEAL if label_col0_teal else b.INK)
        for j, key in enumerate(statuses):
            b.status_cell(cells[1 + j], key)
        comment_cell(cells[7], verdict, comment)
    b.zebra(t)
    b.spacer(doc, 6)
    return t


def build():
    doc = Document()
    b.set_normal_style(doc)
    b.setup_page(doc, landscape=True, margin_cm=1.5)
    b.add_logo_header(doc)
    b.add_footer_pagenum(doc)

    # ---------------- COVER ----------------
    b.cover_page(
        doc,
        doc_kicker="Coverage Comparison \u2014 Media & Entertainment",
        title="Full Coverage Comparison",
        subtitle="A like-for-like, section-by-section comparison of the TMHCC Media & Music Combined wording against four competitor wordings (fifth competitor pending).",
        meta_lines=[
            "Indicative wording comparison against the TMHCC Media & Music Combined wording \u2014 for broking reference only. Always refer to the full policy wordings.",
            "Market-ready subject to final TMHCC legal / underwriting sign-off.   Prepared from the wordings supplied; competitor capacity insurers noted where stated.",
        ],
        wordings=[f"{w['name']} \u2014 {w['full']} ({w['insurer']}); {w['sections']}" for w in d.WORDINGS],
    )

    # ---------------- 2. EXECUTIVE SUMMARY ----------------
    doc.add_page_break()
    b.h1(doc, "Executive summary", num="1")
    b.para(doc, "TMHCC\u2019s Media & Music Combined wording is, structurally, the broadest in the peer group reviewed. It is the only wording that offers all fifteen sections in a single contract, and it is the only wording to provide Loss of Licence, Commercial Legal Expenses, Management Liability and a genuine standalone Cyber-liability section. Two of the four competitors (Liberty and Allianz) are property / business-interruption / liability combined packages and contain no media, professional-indemnity, production-indemnity, cyber, legal-expenses or management-liability cover at all.", align='just')
    b.para(doc, "The meaningful competition therefore sits in the sections the wordings share \u2014 principally Media Liability / Professional Indemnity (where Tysers, via Zurich, and Yutree, via AXA, both offer strong wordings) and the property / BI / liability suite (where Allianz and Liberty are mature property-led packages).", align='just')
    b.h3(doc, "Headline conclusions")
    b.bullet(doc, "TMHCC is broader or unique on at least six rows (Terrorism vs Tysers; Goods in Transit; Loss of Licence; Legal Expenses; Management Liability; standalone Cyber).", bold_lead="TMHCC wins outright:  ")
    b.bullet(doc, "Tysers PI is worldwide, covers patents, funds statutory/criminal defence costs (GBP 1m), and adds a distributors-&-purchasers extension and journalistic source-protection costs. Tysers and Yutree both offer Personal Accident, which TMHCC does not.", bold_lead="Competitors broader:  ")
    b.bullet(doc, "Core media-liability grants (defamation, IP, privacy, professional negligence, defence costs in addition) are broadly equivalent across TMHCC, Tysers and Yutree.", bold_lead="Broadly equivalent:  ")
    b.bullet(doc, "TMHCC\u2019s data-protection defence costs, clean cyber-exclusion carve-out for the liability sections, and favourable fines/penalties write-back are clearer than at least one competitor.", bold_lead="TMHCC clearer:  ")
    b.callout(doc, "Bottom line:", "TMHCC already presents as the most complete media & entertainment package in this peer group. A focused set of Media-Liability enhancements (a Personal-Accident option, a worldwide/US-Canada territory option, distributors & source-protection extensions, and a clarifying criminal-defence cross-reference) would let TMHCC show a tick on virtually every row \u2014 see the separate Gap-Fill / Wording Enhancement Strategy.")

    # ---------------- 3. METHODOLOGY & DOCUMENTS ----------------
    doc.add_page_break()
    b.h1(doc, "Methodology and documents reviewed", num="2")
    b.h3(doc, "Documents reviewed")
    t = b.make_table(doc, ["Wording", "Full title", "Capacity / insurer", "Structure"], [3.6, 9.5, 7.0, 5.9])
    for w in d.WORDINGS:
        cells = t.add_row().cells
        b.text_cell(cells[0], w['name'], size=8.5, bold=True, color=b.TEAL)
        b.text_cell(cells[1], w['full'], size=8.3)
        b.text_cell(cells[2], w['insurer'], size=8.3)
        b.text_cell(cells[3], w['sections'], size=8.3)
    b.zebra(t)
    b.spacer(doc, 6)
    b.h3(doc, "Comparison method")
    for n in d.METHOD_NOTES:
        b.bullet(doc, n)
    b.spacer(doc, 4)
    b.callout(doc, "Reading the matrices:", "\u2713 = cover clearly present \u00b7 \u25D0 = partial / conditional / sub-limited \u00b7 \u2717 = not covered / not identified \u00b7 ? = unclear, requires review \u00b7 \u2014 = not applicable. In the EXCLUSIONS tables only, \u2713 means the exclusion APPLIES, \u25D0 means a partial exclusion or write-back, and \u2717 means the exclusion is absent or carved back.", bg="FBF1DA", border=b.GOLD, lead_color="8A6A1E")

    # ---------------- 4. MARKET COMPARISON AT A GLANCE ----------------
    doc.add_page_break()
    b.h1(doc, "Market comparison at a glance", num="3")
    b.para(doc, "Section-by-section availability across the six wordings (the fifth competitor column is reserved). This is the headline coverage matrix; feature-level detail for the Media Liability / PI section follows in part 4.")
    b.legend_bar(doc)
    render_matrix(doc, "Section of cover", d.SECTION_ROWS)

    # ---------------- 5. FULL COVERAGE MATRIX (feature detail) ----------------
    doc.add_page_break()
    b.h1(doc, "Full coverage matrix \u2014 Media Liability / PI feature detail", num="4")
    b.para(doc, "The Media Liability / Professional Indemnity section (TMHCC S12; Tysers S8; Yutree PI-Media) is the heart of a media wording, so it is compared feature-by-feature below. Liberty and Allianz provide no media/PI cover and are shown as not covered throughout.")
    b.legend_bar(doc)
    render_matrix(doc, "Media Liability / PI feature", d.MEDIA_FEATURES, label_w=5.6, comment_w=11.4)

    # ---------------- 6. COMPETITOR-BY-COMPETITOR ----------------
    doc.add_page_break()
    b.h1(doc, "Competitor-by-competitor analysis", num="5")
    for key in ["tysers", "yutree", "liberty", "allianz"]:
        prof = d.COMPETITOR_PROFILES[key]
        b.h2(doc, prof['title'])
        b.para(doc, prof['shape'], italic=True, color=b.GREY, size=9)
        b.h3(doc, "Where it competes well", color="1B7A3D")
        for s in prof['strengths']:
            b.bullet(doc, s)
        b.h3(doc, "Where TMHCC is broader", color=b.TEAL)
        for g in prof['gaps']:
            b.bullet(doc, g)
        b.spacer(doc, 4)
    b.callout(doc, "Wording E (fifth competitor):", "Pending. A reserved column has been carried through every matrix so the fifth wording can be dropped in without re-formatting.", bg="EEF1F3", border=b.GREY, lead_color=b.GREY)

    # ---------------- 7. TMHCC STRENGTHS ----------------
    doc.add_page_break()
    b.h1(doc, "TMHCC strengths", num="6")
    b.para(doc, "Where the TMHCC base wording is demonstrably stronger, broader or clearer than the competitor wordings reviewed:")
    t = b.make_table(doc, ["Strength", "Why it matters"], [7.0, 19.0])
    for title, why in d.TMHCC_STRENGTHS:
        cells = t.add_row().cells
        b.text_cell(cells[0], title, size=8.6, bold=True, color=b.TEAL)
        b.text_cell(cells[1], why, size=8.6)
    b.zebra(t)

    # ---------------- 8. AREAS WHERE COMPETITORS ARE BROADER ----------------
    doc.add_page_break()
    b.h1(doc, "Areas where competitors are broader", num="7")
    b.para(doc, "Stated objectively \u2014 where a competitor genuinely provides broader, clearer or materially useful cover. Each item is carried into the separate Gap-Fill / Wording Enhancement Strategy.")
    t = b.make_table(doc, ["Area (competitor)", "What the competitor does / TMHCC position"], [7.0, 19.0])
    for title, why in d.COMP_BROADER:
        cells = t.add_row().cells
        b.text_cell(cells[0], title, size=8.6, bold=True, color="A4233B")
        b.text_cell(cells[1], why, size=8.6)
    b.zebra(t)

    # ---------------- 9. EXCLUSIONS COMPARISON ----------------
    doc.add_page_break()
    b.h1(doc, "Exclusions comparison", num="8")
    b.h2(doc, "General exclusions (all sections)")
    b.callout(doc, "Note:", "In this table \u2713 = the exclusion APPLIES, \u25D0 = partial / write-back present, \u2717 = exclusion absent or carved back, ? = requires review.", bg="FBF1DA", border=b.GOLD, lead_color="8A6A1E")
    render_matrix(doc, "General exclusion", [(r[0], r[1], None, r[2]) for r in d.GENERAL_EXCL], label_w=5.6, comment_w=11.4)
    b.h2(doc, "Media Liability / PI-specific exclusions")
    b.para(doc, "Liberty and Allianz have no PI section and are shown as not applicable.")
    render_matrix(doc, "PI exclusion", [(r[0], r[1], None, r[2]) for r in d.PI_EXCL], label_w=5.6, comment_w=11.4)
    b.callout(doc, "Exclusions \u2014 key takeaways:", "TMHCC is sensibly MORE protected on communicable disease, patents and US/Canada jurisdiction. The genuine write-back opportunities a competitor offers are: asbestos professional-duty (Tysers), pollution negligent-advice (Yutree) and a rated worldwide/US-Canada PI option (Tysers). TMHCC is clearer where its cyber/date exclusion is dis-applied to the liability sections, avoiding a \u2018silent-cyber strip\u2019 of PI.")

    # ---------------- 10. CONDITIONS / WARRANTIES / CLAIMS ----------------
    doc.add_page_break()
    b.h1(doc, "Conditions, warranties and claims obligations", num="9")
    headers = ["Condition / obligation", "TMHCC", "Tysers (Zurich)", "Yutree (AXA)", "Liberty", "Allianz"]
    t = b.make_table(doc, headers, [4.6, 5.0, 4.3, 4.3, 4.0, 3.8])
    for row in d.CONDITIONS:
        cells = t.add_row().cells
        b.text_cell(cells[0], row[0], size=8.2, bold=True, color=b.TEAL)
        for j in range(5):
            b.text_cell(cells[1 + j], row[1 + j], size=7.9)
    b.zebra(t)
    b.callout(doc, "Claims obligations:", "TMHCC\u2019s claim-notification is an EXPRESS condition precedent to liability \u2014 among the most protective drafting in the peer group, and a point to retain (not soften). Yutree\u2019s King\u2019s Counsel dispute clause is a drafting-clarity advantage TMHCC could adopt.")

    # ---------------- 11. LIMITS / SUB-LIMITS / EXCESS ----------------
    doc.add_page_break()
    b.h1(doc, "Limits, sub-limits and excess comparison", num="10")
    b.para(doc, "Primary limits and excesses are schedule-driven and were not supplied; the table below compares the principal SUB-LIMITS within the Media Liability / PI sections (Liberty and Allianz have no PI section).")
    headers = ["Media / PI sub-limit", "TMHCC", "Tysers (Zurich)", "Yutree (AXA)", "Liberty", "Allianz"]
    t = b.make_table(doc, headers, [5.4, 4.4, 4.6, 4.6, 3.5, 3.5])
    for row in d.SUBLIMITS:
        cells = t.add_row().cells
        b.text_cell(cells[0], row[0], size=8.3, bold=True, color=b.TEAL)
        for j in range(5):
            b.text_cell(cells[1 + j], row[1 + j], size=8.0, align='center')
    b.zebra(t)
    b.callout(doc, "Excess:", "TMHCC and the competitors apply a per-claim excess per the Schedule; both Tysers and TMHCC pay defence costs in addition to the limit (Tysers states the excess does not apply to defence costs). Confirm final limits/excesses against each Schedule.")

    # ---------------- 12. LEGAL / UNDERWRITING REVIEW NOTES ----------------
    doc.add_page_break()
    b.h1(doc, "Legal / underwriting review notes", num="11")
    b.para(doc, "Items flagged \u2018requires review\u2019 in the matrices, plus points for TMHCC legal / underwriting sign-off:")
    notes = [
        "Confirm whether TMHCC offers Personal Accident / Business Travel by endorsement (no standalone section identified).",
        "Allianz references a \u2018Personal Accident\u2019 section in its exclusions that is not in the provided cover list \u2014 confirm against the Allianz schedule.",
        "Confirm communicable-disease exclusion positions in the Yutree, Liberty and Allianz wordings (TMHCC carries a broad exclusion).",
        "Confirm Tysers and Yutree positions on fines/penalties and insolvency exclusions within PI (marked \u2018review\u2019).",
        "Confirm Increased Cost of Working / Book Debts and computer-breakdown are expressly addressed within TMHCC S3 and S1/S2 respectively.",
        "Primary limits, excesses and any \u2018insured/not insured\u2019 toggles are schedule-driven and must be confirmed against each Schedule.",
        "This comparison is indicative and for broking reference; it is subject to final TMHCC legal / underwriting sign-off before market use.",
    ]
    for n in notes:
        b.bullet(doc, n)

    # ---------------- 13. APPENDIX: CLAUSE MAPPING ----------------
    doc.add_page_break()
    b.h1(doc, "Appendix \u2014 detailed clause mapping", num="12")
    b.para(doc, "Cross-reference of comparable cover by FUNCTION across the wordings (section names are the actual headings used in each policy).")
    mapping = [
        ("Premises material damage", "S1 Business \u2018All Risks\u2019", "S2 Property Damage", "Property damage \u2013 All risks", "S1 Material Damage", "S1 Property Damage"),
        ("Production / entertainment equipment", "S2 Property & Equipment", "S1 Production Property", "Production property", "Within S1 / Schedule", "S6 Specified All Risks"),
        ("Business interruption", "S3 Business Interruption", "S3 BI \u2018All Risks\u2019", "BI \u2013 All risks", "S2 Business Interruption", "S2 BI / S2a ICOW / S3 Book Debts"),
        ("Terrorism", "S4 Terrorism", "Excluded", "Terrorism section", "S3 Terrorism", "Terrorism"),
        ("Employers\u2019 liability", "S5 Employers\u2019 Liability", "S6 Employers\u2019 Liability", "Employers liability", "S5 Employer\u2019s Liability", "S7 Employers\u2019 Liability"),
        ("Public & products liability", "S6 Public / S7 Products", "S7 Public & Products", "Public & products liability", "S6 Public & Products", "S8 Public & Products"),
        ("Money", "S8 Money", "S4 Money", "Money & PA assault", "S4 Money", "S4 Money"),
        ("Goods in transit", "S9 Goods in Transit", "\u2014", "Goods in transit", "\u2014", "S5 Own Goods in Transit"),
        ("Loss of licence", "S10 Loss of Licence", "\u2014", "\u2014", "\u2014", "\u2014"),
        ("Production indemnity", "S11 Production Indemnity", "S5 Multimedia & Producers Indemnity", "Production indemnity \u2013 All risks", "\u2014", "\u2014"),
        ("Media liability / PI / E&O", "S12 Media Liability", "S8 Professional Indemnity/E&O", "PI \u2013 Media; PI \u2013 Events", "\u2014", "\u2014"),
        ("Legal expenses", "S13 Commercial Legal Expenses", "\u2014", "\u2014", "\u2014", "\u2014"),
        ("Management liability", "S14 Management Liability", "\u2014", "\u2014", "\u2014", "\u2014"),
        ("Cyber liability (standalone)", "S15 CyberGuard", "\u2014", "Computer breakdown (first-party only)", "\u2014", "\u2014"),
        ("Personal accident / travel", "\u2014", "S9 PA & Business Travel", "PA assault (within Money/PA)", "\u2014", "Referenced \u2013 review"),
    ]
    headers = ["Cover (by function)", "TMHCC", "Tysers (Zurich)", "Yutree (AXA)", "Liberty", "Allianz"]
    t = b.make_table(doc, headers, [4.6, 4.8, 4.8, 4.6, 3.6, 3.6])
    for row in mapping:
        cells = t.add_row().cells
        b.text_cell(cells[0], row[0], size=8.0, bold=True, color=b.TEAL)
        for j in range(5):
            b.text_cell(cells[1 + j], row[1 + j], size=7.8)
    b.zebra(t)
    b.spacer(doc, 6)
    b.para(doc, "End of Full Coverage Comparison. A companion document \u2014 TMHCC Gap-Fill / Wording Enhancement Strategy \u2014 sets out the recommended enhancements arising from this comparison.", italic=True, color=b.GREY, size=8.5)

    out = "/app/work/compare/out/TMHCC_Media_Coverage_Comparison_FULL.docx"
    b.save_doc(doc, out)
    print("Saved:", out)
    return out


if __name__ == "__main__":
    build()
