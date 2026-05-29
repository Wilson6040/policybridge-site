"""
NEW_0526 wording — two-column TOC variant.

Same corrections as the live-TOC version (Section 16 removed, stale IT-section
references removed, CyberGuard restyled, footer page numbers) BUT the table of
contents is a TWO-COLUMN, hyperlinked list with accurate page numbers
(right-aligned with dot leaders).  Page numbers are computed from the actual
rendered document via the UNO page cursor, so they are correct and consistent.

Run order:
  python3 build_2col.py            # builds _2col_nonum.docx (placeholder numbers)
  /usr/bin/python3 uno_pagenums.py _2col_nonum.docx tocpg_ pg2col.json
  python3 build_2col.py --inject   # writes page numbers, saves FINAL docx
"""
import sys, copy, zipfile, json
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = '/app/work/source/NEW_0526.docx'
NONUM = '/app/work/output/_2col_nonum.docx'
FINAL = '/app/work/output/TMHCC_Media_Combined_0526_FINAL_amended.docx'
PGJSON = '/app/work/pg2col.json'
COLW = 4862               # column width (twips)
COLBREAK_AT = 13          # entry index that starts column 2

ENTRIES = [
    (23,  "Policy Contents", 1),
    (74,  "Introduction", 1),
    (82,  "Notices", 1),
    (152, "Insuring Agreement", 1),
    (159, "General Policy Definitions", 1),
    (309, "General Policy Conditions", 1),
    (450, "Policy Protection and Maintenance Conditions", 1),
    (529, "Policy Claims Conditions", 1),
    (610, "Policy Exclusions", 1),
    (698, "Section 1: Business \u201cAll Risks\u201d Section", 1),
    (978, "Section 2: Property & Equipment \u201cAll Risks\u201d Section", 1),
    (1205, "Section 3: Business Interruption \u201cAll Risks\u201d Section", 1),
    (1520, "Section 4: Terrorism Section", 1),
    (1612, "Section 5: Employers\u2019 Liability Section", 1),
    (1671, "Section 6: Public Liability Section", 1),
    (1846, "Section 7: Products Liability Section", 1),
    (1914, "Section 8: Money Section", 1),
    (1915, "Sub-Section 1 \u2013 Money", 2),
    (1962, "Sub-Section 2 \u2013 Personal Assault", 2),
    (1995, "Section 9: Goods in Transit Section", 1),
    (2063, "Section 10: Loss of Licence Section", 1),
    (2115, "Section 11: Production Indemnity \u201cAll Risks\u201d Section", 1),
    (2222, "Section 12: Media Liability Section", 1),
    (2466, "Section 13: Commercial Legal Expenses Section", 1),
    (2784, "Section 14: Management Liability Insurance Section", 1),
    (3156, "Section 15: Cyber Liability Section", 1),
]

SECT16_FIX = {
    156: ("Sections 13, 14, 15 or 16", "Sections 13, 14 or 15"),
    310: ("Sections 15 and 16", "Sections 14 and 15"),
    611: ("Sections 13,15 and 16", "Sections 13, 14 and 15"),
}
IT_FIX = {
    233: ("the Information Technology Section", "the Business \u201cAll Risks\u201d Section"),
    384: (" or the Information Technology Section", ""),
    418: (", Information Technology \u201cAll Risks\u201d", ""),
    451: (", Information Technology and Money Sections", " and Money Sections"),
    473: (" and/or the Information Technology Section", ""),
    646: (", Information Technology, Money", ", Money"),
    656: (", Information Technology and Production Indemnity", " and Production Indemnity"),
}
IT_DELETE = [1567]
CYBER_SUBTITLE = [3156]
CYBER_H1 = [3160, 3244, 3649, 3824, 3898, 3940, 3951, 4005, 4017]
CYBER_H2 = [3164, 3190, 3736, 3757, 3767, 3784, 3828, 3856, 3878, 3886, 3892, 3895,
            3925, 3926, 3929, 3941, 3948, 3953, 3958, 3962, 3968, 3971, 3974, 3982,
            3985, 3990, 3993, 4000, 4002, 4009, 3899, 3910, 3915]
CYBER_EXCLUDE = {3157, 3646}
CYBER_RANGE = (3156, 4037)


def rfonts(rpr):
    rf = OxmlElement('w:rFonts')
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rf.set(qn(a), 'Arial')
    rpr.append(rf)


def run(text, sz='19', color='000000', bold=False):
    r = OxmlElement('w:r'); rpr = OxmlElement('w:rPr'); rfonts(rpr)
    if bold:
        rpr.append(OxmlElement('w:b'))
    s = OxmlElement('w:sz'); s.set(qn('w:val'), sz); rpr.append(s)
    c = OxmlElement('w:color'); c.set(qn('w:val'), color); rpr.append(c)
    r.append(rpr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text; r.append(t)
    return r


def make_entry(text, anchor, level, colbreak):
    p = OxmlElement('w:p'); pPr = OxmlElement('w:pPr')
    sp = OxmlElement('w:spacing'); sp.set(qn('w:after'), '40'); sp.set(qn('w:line'), '264'); sp.set(qn('w:lineRule'), 'auto'); pPr.append(sp)
    if level == 2:
        ind = OxmlElement('w:ind'); ind.set(qn('w:left'), '227'); pPr.append(ind)
    tabs = OxmlElement('w:tabs'); tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right'); tab.set(qn('w:pos'), str(COLW)); tab.set(qn('w:leader'), 'dot')
    tabs.append(tab); pPr.append(tabs)
    p.append(pPr)
    if colbreak:
        br_r = OxmlElement('w:r'); br = OxmlElement('w:br'); br.set(qn('w:type'), 'column'); br_r.append(br); p.append(br_r)
    hl = OxmlElement('w:hyperlink'); hl.set(qn('w:anchor'), anchor); hl.set(qn('w:history'), '1')
    hl.append(run(text, bold=(level == 1)))
    p.append(hl)
    tr = OxmlElement('w:r'); tr.append(OxmlElement('w:tab')); p.append(tr)
    p.append(run('', bold=(level == 1)))   # placeholder page-number run (last run)
    return p


def add_bookmark(p_el, name, bid):
    bs = OxmlElement('w:bookmarkStart'); bs.set(qn('w:id'), str(bid)); bs.set(qn('w:name'), name)
    be = OxmlElement('w:bookmarkEnd'); be.set(qn('w:id'), str(bid))
    pPr = p_el.find(qn('w:pPr'))
    (pPr.addnext(bs) if pPr is not None else p_el.insert(0, bs))
    bs.addnext(be)


def para_text(p_el):
    return ''.join(t.text or '' for t in p_el.iter(qn('w:t')))


def set_single_text(p_el, new_text):
    runs = p_el.findall(qn('w:r')); done = False
    for r in runs:
        ts = r.findall(qn('w:t'))
        if not done and ts:
            ts[0].text = new_text; ts[0].set(qn('xml:space'), 'preserve')
            for x in ts[1:]:
                x.text = ''
            done = True
        else:
            for t in r.findall(qn('w:t')):
                t.text = ''


def replace_in_para(p_el, old, new):
    txt = para_text(p_el)
    if old in txt:
        set_single_text(p_el, txt.replace(old, new, 1)); return True
    return False


def set_pagebreak_before(p_el):
    pPr = p_el.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr'); p_el.insert(0, pPr)
    if pPr.find(qn('w:pageBreakBefore')) is None:
        pPr.insert(0, OxmlElement('w:pageBreakBefore'))


# ===================== inject mode =====================
if '--inject' in sys.argv:
    pages = json.load(open(PGJSON))
    doc = Document(NONUM)
    paras = doc.paragraphs
    # locate "Policy Contents" heading, the 26 entries follow it
    pc_idx = next(i for i, p in enumerate(paras) if p.text.strip() == 'Policy Contents'
                  and p.style.name.startswith('Heading'))
    for i, (tgt, text, level) in enumerate(ENTRIES):
        entry = paras[pc_idx + 1 + i]
        pg = pages.get(f"tocpg_{i}")
        runs = entry._p.findall(qn('w:r'))
        # last run is the page-number placeholder
        last = runs[-1]
        for t in last.findall(qn('w:t')):
            t.text = str(pg) if pg else ''
            t.set(qn('xml:space'), 'preserve')
    doc.save('/app/work/output/_2col_withnum.docx')
    # footer page numbers
    PAGE_PARA = ('<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                 '<w:pPr><w:jc w:val="center"/></w:pPr>'
                 '<w:r><w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>'
                 '<w:color w:val="595959"/><w:sz w:val="16"/></w:rPr><w:fldChar w:fldCharType="begin"/></w:r>'
                 '<w:r><w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>'
                 '<w:color w:val="595959"/><w:sz w:val="16"/></w:rPr>'
                 '<w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>'
                 '<w:r><w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>'
                 '<w:color w:val="595959"/><w:sz w:val="16"/></w:rPr><w:fldChar w:fldCharType="end"/></w:r></w:p>')
    FOOTERS = {'word/footer1.xml', 'word/footer2.xml', 'word/footer3.xml'}
    zin = zipfile.ZipFile('/app/work/output/_2col_withnum.docx', 'r')
    zout = zipfile.ZipFile(FINAL, 'w', zipfile.ZIP_DEFLATED)
    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename in FOOTERS:
            data = data.decode('utf-8').replace('</w:ftr>', PAGE_PARA + '</w:ftr>').encode('utf-8')
        zout.writestr(item, data)
    zin.close(); zout.close()
    print("FINAL (2-column TOC) saved:", FINAL)
    sys.exit(0)

# ===================== build mode =====================
doc = Document(SRC)
paras = doc.paragraphs
ref = {}
needed = set(i for i, _, _ in ENTRIES) | set(SECT16_FIX) | set(IT_FIX) | set(IT_DELETE)
needed |= set(CYBER_SUBTITLE) | set(CYBER_H1) | set(CYBER_H2) | {22, 23}
needed |= set(range(24, 50)) | set(range(*CYBER_RANGE))
for i in needed:
    if i < len(paras):
        ref[i] = paras[i]

# content fixes
for idx, (old, new) in SECT16_FIX.items():
    replace_in_para(ref[idx]._p, old, new)
for idx, (old, new) in IT_FIX.items():
    replace_in_para(ref[idx]._p, old, new)

# cyber restyle
for idx in CYBER_SUBTITLE:
    ref[idx].style = doc.styles['Subtitle']
for idx in CYBER_H1:
    if idx in ref:
        ref[idx].style = doc.styles['Heading 1']
for idx in CYBER_H2:
    if idx in ref:
        ref[idx].style = doc.styles['Heading 2']
styled = set(CYBER_SUBTITLE) | set(CYBER_H1) | set(CYBER_H2) | CYBER_EXCLUDE
for i in range(*CYBER_RANGE):
    if i not in ref:
        continue
    t = ref[i].text.strip()
    if t and len(t) <= 60 and t[-1:] not in '.;,:' and len(t) > 2 and not t[0].islower() and i not in styled:
        ref[i].style = doc.styles['Heading 3']

# bookmarks at the 26 targets
bid = 6000
for i, (tgt, text, level) in enumerate(ENTRIES):
    add_bookmark(ref[tgt]._p, f"tocpg_{i}", bid); bid += 1

# TOC: page break on Policy Contents, then build 26 entries after it, delete old 24-49
set_pagebreak_before(ref[23]._p)
del_paras = [ref[i]._p for i in range(24, 50)]
anchor_el = ref[23]._p
for i, (tgt, text, level) in enumerate(ENTRIES):
    ep = make_entry(text, f"tocpg_{i}", level, colbreak=(i == COLBREAK_AT))
    anchor_el.addnext(ep)
    anchor_el = ep
for el in del_paras:
    el.getparent().remove(el)
for idx in IT_DELETE:
    el = ref[idx]._p
    el.getparent().remove(el)

doc.save(NONUM)
print("built _2col_nonum.docx (placeholder page numbers)")
