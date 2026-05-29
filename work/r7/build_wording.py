"""
Round 7 — TMHCC wording tracked-change builder.
Produces:
  baseline_clean.docx          (existing tracked changes accepted; internal baseline)
  ..._TrackedChanges.docx      (baseline + NEW tracked changes for Tasks 1,2,4,11)
  ..._Clean_Final.docx         (accept-all of the new tracked changes)
"""
import copy, datetime
from docx import Document
from lxml import etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
def qn(t): return f"{{{W}}}{t}"
def ln(el): return etree.QName(el).localname

AUTHOR = "TMHCC Wording Review"
DATE = "2025-07-01T00:00:00Z"
_id = [9000]
def nid():
    _id[0] += 1
    return str(_id[0])

SRC = "/app/work/r7/source_uploaded.docx"

# ---------- low-level builders ----------
def make_t(text):
    t = etree.SubElement(etree.Element(qn("r")), qn("t"))  # placeholder
    return None

def new_run(text, bold=False, rpr_src=None):
    r = etree.Element(qn("r"))
    rpr = None
    if rpr_src is not None:
        # copy an existing rPr element
        src = rpr_src.find(qn("rPr"))
        if src is not None:
            rpr = copy.deepcopy(src)
    if bold:
        if rpr is None:
            rpr = etree.Element(qn("rPr"))
        if rpr.find(qn("b")) is None:
            rpr.append(etree.Element(qn("b")))
    if rpr is not None:
        r.append(rpr)
    t = etree.SubElement(r, qn("t"))
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return r

def wrap_ins(children):
    ins = etree.Element(qn("ins"))
    ins.set(qn("id"), nid()); ins.set(qn("author"), AUTHOR); ins.set(qn("date"), DATE)
    for c in children:
        ins.append(c)
    return ins

def wrap_del(runs):
    """Wrap given run elements into a w:del, converting w:t -> w:delText."""
    delel = etree.Element(qn("del"))
    delel.set(qn("id"), nid()); delel.set(qn("author"), AUTHOR); delel.set(qn("date"), DATE)
    for r in runs:
        for t in r.findall(qn("t")):
            t.tag = qn("delText")
        delel.append(r)
    return delel

def mark_para_inserted(p):
    """Mark a paragraph's end-mark as an inserted revision."""
    pPr = p.find(qn("pPr"))
    if pPr is None:
        pPr = etree.Element(qn("pPr")); p.insert(0, pPr)
    rPr = pPr.find(qn("rPr"))
    if rPr is None:
        rPr = etree.SubElement(pPr, qn("rPr"))
    ins = etree.SubElement(rPr, qn("ins"))
    ins.set(qn("id"), nid()); ins.set(qn("author"), AUTHOR); ins.set(qn("date"), DATE)

def new_para(style, run_segments):
    """run_segments: list of (text, bold). Returns a w:p with inserted runs + inserted para mark."""
    p = etree.Element(qn("p"))
    pPr = etree.SubElement(p, qn("pPr"))
    ps = etree.SubElement(pPr, qn("pStyle")); ps.set(qn("val"), style)
    runs = [new_run(txt, bold) for (txt, bold) in run_segments]
    p.append(wrap_ins(runs))
    mark_para_inserted(p)
    return p

# ---------- accept helpers ----------
def accept_all(body):
    root = body.getroottree()
    # 1. unwrap w:ins (move children to parent)
    for ins in body.iter(qn("ins")):
        parent = ins.getparent()
        if parent is None:
            continue
        # paragraph-mark ins (inside rPr) -> just drop it
        if ln(parent) == "rPr":
            parent.remove(ins)
        else:
            idx = list(parent).index(ins)
            for c in reversed(list(ins)):
                parent.insert(idx, c)
            parent.remove(ins)
    # 2. remove w:del entirely
    for d in list(body.iter(qn("del"))):
        if d.getparent() is not None:
            d.getparent().remove(d)
    # 3. remove rPrChange / pPrChange
    for tag in ("rPrChange", "pPrChange"):
        for e in list(body.iter(qn(tag))):
            if e.getparent() is not None:
                e.getparent().remove(e)

def accept_existing(doc):
    """Accept the few pre-existing tracked changes (rPrChange etc) to form baseline."""
    body = doc.element.body
    accept_all(body)

# ---------- find helpers ----------
def find_para(doc, predicate):
    for p in doc.paragraphs:
        if predicate(p):
            return p
    return None

def runs_of(p_el):
    return [c for c in p_el if ln(c) == "r"]

# ====================================================================
# STEP 1 — baseline (accept existing)
# ====================================================================
doc = Document(SRC)
accept_existing(doc)
doc.save("/app/work/r7/baseline_clean.docx")
print("baseline_clean.docx saved")

# ====================================================================
# STEP 2 — apply NEW tracked changes on the baseline
# ====================================================================
doc = Document("/app/work/r7/baseline_clean.docx")

# ---- TASK 1: replace Admission of Liability CP sentence [507] ----
adm_txt = "It is a condition precedent to liability under this Policy that no admission of liability, promise, payment, compensation, negotiation or settlement of any claim shall be made or given without the Insurer"
p507 = find_para(doc, lambda p: (p.text or "").strip().startswith(adm_txt[:60]) and "no admission of liability" in p.text)
assert p507 is not None, "Task1 para not found"
pe = p507._p
existing = runs_of(pe)
delel = wrap_del(existing)  # this moves runs out of pe into delel (and converts t->delText)
# build new sentence runs
seg = [
    ("The ", False), ("Insured", True),
    (" shall not, without the ", False), ("Insurer\u2019s", True),
    (" written consent, make any admission of liability, promise, payment, compensation, offer, negotiation or settlement of any claim. The ", False),
    ("Insurer", True),
    (" shall not be entitled to rely on any breach of this condition to reduce or refuse a claim except to the extent that the ", False),
    ("Insurer\u2019s", True),
    (" position has been prejudiced by the breach.", False),
]
insel = wrap_ins([new_run(t, b) for (t, b) in seg])
pe.append(delel)
pe.append(insel)
print("Task1 applied")

# ---- TASK 2: insert Proof of Ownership and Value after police-CP body [546] ----
police_key = "if theft or attempted theft or"
p546 = find_para(doc, lambda p: "Condition Precedent to liability" in (p.text or "") and police_key in (p.text or ""))
assert p546 is not None, "Task2 anchor (police CP) not found"
anchor = p546._p
head_p = new_para("Heading3", [("Proof of Ownership and Value", False)])
body_segs = [
    ("Following any claim for theft or loss, the ", False), ("Insurer", True),
    (" may request reasonable evidence of ownership and value in respect of any high value item or any item insured on an ", False),
    ("Agreed Value", True),
    (" basis (including but not limited to purchase invoices, professional valuations, photographs, serial numbers or other documentary evidence). The ", False),
    ("Insured", True),
    (" shall provide such evidence within a reasonable period following the ", False),
    ("Insurer\u2019s", True),
    (" request. The ", False), ("Insurer", True),
    (" shall not be obliged to pay any claim, or part of a claim, for an affected item where the ", False),
    ("Insured", True),
    (" fails to provide the evidence reasonably requested.", False),
]
body_p = new_para("BodyText", body_segs)
anchor.addnext(body_p)   # body after anchor
head_p_el = head_p
anchor.addnext(head_p_el)  # heading after anchor (so order: anchor, heading, body)
print("Task2 applied")

# ---- TASK 4 (Gap 8): label Accounts Receivable as Book Debts [1205],[1253] ----
ar_paras = [p for p in doc.paragraphs if (p.text or "").strip() == "Accounts Receivable" and (p.style and p.style.name == "Heading 3")]
assert len(ar_paras) >= 2, f"Task4: expected >=2 Accounts Receivable headings, found {len(ar_paras)}"
for p in ar_paras:
    last_run = runs_of(p._p)[-1]
    insr = wrap_ins([new_run(" (Book Debts)", False, rpr_src=last_run)])
    last_run.addnext(insr)
print(f"Task4 applied to {len(ar_paras)} headings")

# ---- TASK 11: standardise Section 15 label ----
# body [3203]: insert " (Cyber Liability)" after CyberGuard run
pbody = find_para(doc, lambda p: (p.text or "").strip() == "Section 15 - CyberGuard\u2122")
assert pbody is not None, "Task11 body heading not found"
for r in runs_of(pbody._p):
    ts = r.findall(qn("t"))
    if ts and "CyberGuard" in (ts[-1].text or ""):
        insr = wrap_ins([new_run(" (Cyber Liability)", False, rpr_src=r)])
        r.addnext(insr)
        break
print("Task11 body applied")

# contents [48]: replace ' Cyber Liability Section' inside hyperlink with ' CyberGuard (Cyber Liability)'
ptoc = find_para(doc, lambda p: (p.text or "").strip().startswith("Section 15:") and "Cyber Liability Section" in (p.text or ""))
assert ptoc is not None, "Task11 contents entry not found"
hl = ptoc._p.find(qn("hyperlink"))
assert hl is not None, "Task11 contents hyperlink not found"
target_run = None
for r in [c for c in hl if ln(c) == "r"]:
    ts = r.findall(qn("t"))
    if ts and "Cyber Liability Section" in (ts[-1].text or ""):
        target_run = r
        break
assert target_run is not None, "Task11 contents target run not found"
new_r = new_run(" CyberGuard\u2122 (Cyber Liability)", False, rpr_src=target_run)
delel = wrap_del([target_run])  # moves target_run into del, converts t->delText
insel = wrap_ins([new_r])
# place del then ins where target_run was (target_run already detached into delel)
# insert at end of hyperlink's run sequence (before any trailing) — simplest: append in order at the position
# find index: append after the first run ("Section 15:")
first_run = [c for c in hl if ln(c) == "r"][0]
first_run.addnext(insel)
first_run.addnext(delel)
print("Task11 contents applied")

doc.save("/app/work/r7/TMHCC_Media_Combined_0526_FINAL_TrackedChanges.docx")
print("TrackedChanges.docx saved")

# ====================================================================
# STEP 3 — clean final (accept all new tracked changes)
# ====================================================================
doc2 = Document("/app/work/r7/TMHCC_Media_Combined_0526_FINAL_TrackedChanges.docx")
accept_all(doc2.element.body)
doc2.save("/app/work/r7/TMHCC_Media_Combined_0526_FINAL_Clean.docx")
print("Clean_Final.docx saved")
